import asyncio
import base64
import calendar
from collections import defaultdict
import io
import json
import math
import os
import random
import re
import sqlite3
import string
import tempfile
import time
import warnings
import webbrowser
from datetime import date, datetime, timedelta

import flet as ft
import fitz
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx2pdf import convert

warnings.filterwarnings("ignore", message=".*OpenSSL.*LibreSSL.*", module="urllib3")
warnings.filterwarnings("ignore", message=".*urllib3 v2 only supports OpenSSL.*")

import mntu_app.flet_compat

from mntu_app.config import (
    DEFAULT_SCHEDULE_PATH,
    LIB_BASE_URL,
    MEET_LINKS,
    NEWS_URL,
    REQUEST_TIMEOUT,
    THEMES,
    UKR_MONTH_NAMES,
)
from mntu_app.database import (
    clear_schedule_cache,
    clear_today_notifications,
    delete_announcements,
    get_all_announcements,
    get_shown_notifications_for_today,
    get_unread_announcements,
    init_db,
    load_schedule_from_db,
    load_settings,
    mark_announcements_seen,
    mark_notification_shown,
    save_announcement,
    save_custom_accent,
    save_schedule_to_db,
    save_settings_to_db,
    save_telegram_settings,
    save_telegram_auth_profile,
    save_theme_to_db,
    save_user_role_to_db,
    create_role_change_request,
    get_user_role_by_email,
    set_user_role_by_email,
    upsert_user_profile,
    DEFAULT_ADMIN_EMAIL,
)
from mntu_app.flet_compat import (
    bind_dropdown_change as _bind_dropdown_change,
    border_all as _border_all,
    dropdown_event_kw as _dropdown_event_kw,
    dropdown_menu_height_kw as _dropdown_menu_height_kw,
    dropdown_value_from_event as _dropdown_value_from_event,
    padding_all as _padding_all,
    padding_only as _padding_only,
    padding_symmetric as _padding_symmetric,
)
from mntu_app.file_media import (
    docx_to_pdf_images,
    get_full_file_content,
    get_last_2_pages_content,
    pdf_pages_to_base64,
)
from mntu_app.istu_library import File, get_parser, reset_parser
from mntu_app.news_parser import parse_istu_news
from mntu_app.notifications import (
    show_class_notification,
    show_pending_announcements,
    show_snack,
    show_system_notification,
)
from mntu_app.schedule_dates import (
    _is_synthetic_no_classes_entry,
    extract_date_from_text,
    extract_teacher_from_text,
    format_schedule_date_short,
    format_week_range_human,
    parse_schedule_date_str,
    schedule_has_entries,
)
from mntu_app.telegram_bridge import TelegramBridge
from mntu_app.theme import (
    accent_text_color,
    get_meet_link,
    hex_to_rgb,
    hsl_to_hex,
    hsl_to_rgb,
    rgb_to_hex,
    rgb_to_hsl,
)
def main(page: ft.Page):
    init_db()

    page.title = "МНТУ Помічник"
    page.theme_mode = ft.ThemeMode.DARK
    page.window.width = 990
    page.window.height = 730
    page.window.resizable = True
    page.window.maximizable = True
    page.padding = 0
    page.spacing = 0

    settings = load_settings()
    today = date.today()
    theme_key = settings.get("theme", "green")
    colors = dict(THEMES.get(theme_key, THEMES["green"]))
    if theme_key == "custom" and settings.get("custom_accent"):
        colors["ACCENT"] = settings["custom_accent"] if settings["custom_accent"].startswith("#") else "#" + settings["custom_accent"]
    colors["ACCENT_TEXT"] = accent_text_color(colors["ACCENT"])
    page.bgcolor = colors["BG_DARK"]

    schedule_data, group_options, week_options, week_dates = load_schedule_from_db()
    if not schedule_has_entries(schedule_data):
        schedule_data, group_options, week_options, week_dates = {}, [], [], {}

    def _canonical_group_codes(text: str) -> list[str]:
        src = (text or "").strip()
        if not src:
            return []


        matches = re.findall(r"(?<!\w)[А-ЯІЇЄҐA-Z]{1,6}\s*-?\s*\d{1,3}(?:\s*-\s*\d{1,2})?(?!\w)", src.upper())
        out = []
        for m in matches:
            compact = re.sub(r"\s+", "", m)
            compact = compact.replace("–", "-").replace("—", "-")
            compact = re.sub(r"^([А-ЯІЇЄҐA-Z]{1,5})(\d)", r"\1-\2", compact)
            out.append(compact)
        return out

    def normalize_group_names(raw_groups: list[str]) -> list[str]:
        out: list[str] = []
        seen: set[str] = set()
        for item in raw_groups or []:
            txt = (item or "").strip()
            if not txt:
                continue
            parts = [p.strip() for p in re.split(r"\s*,\s*", txt) if p.strip()]
            if not parts:
                parts = [txt]
            for p in parts:
                codes = _canonical_group_codes(p)
                candidates = codes[:]
                if not candidates:

                    fallback = re.sub(r"\([^)]*\)", "", p).strip()
                    if len(fallback) <= 40 and re.search(r"[A-Za-zА-Яа-яІіЇїЄєҐґ]", fallback) and re.search(r"\d", fallback):
                        candidates = [fallback.upper().replace(" ", "")]
                for code in candidates:
                    if code not in seen:
                        seen.add(code)
                        out.append(code)
        return out

    state = {
        "theme": settings.get("theme", "green"),
        "custom_accent": settings.get("custom_accent", "") or colors["ACCENT"].lstrip("#"),
        "parsed_schedule_per_group": schedule_data if schedule_data else {},
        "notify_group": settings["notify_group"],
        "group_options": normalize_group_names(group_options if group_options else []),
        "week_options": week_options if week_options else [],
        "selected_group": "",
        "selected_week": "",
        "schedule_root_url": LIB_BASE_URL + DEFAULT_SCHEDULE_PATH,
        "schedule_url": LIB_BASE_URL + DEFAULT_SCHEDULE_PATH,
        "schedule_stack": [],
        "selected_index": 0,
        "calendar_to_schedule_date": None,
        "schedule_files": [],
        "schedule_loading": False,
        "schedule_error": "",
        "week_dates": week_dates if week_dates else {},
        "viewing_images": None,
        "notifications_enabled": settings["notifications_enabled"],
        "notify_minutes": settings["notify_minutes"],
        "user_role": settings.get("user_role", "Студент"),
        "notified_entries": get_shown_notifications_for_today(),
        "visible_year": today.year,
        "visible_month": today.month,
        "selected_date": today,
        "news_items": [],
        "telegram_enabled": settings.get("telegram_enabled", False),
        "telegram_bot_token": settings.get("telegram_bot_token", ""),
        "telegram_chat_id": settings.get("telegram_chat_id", ""),
        "auth_required": settings.get("auth_required", True),
        "tg_auth_verified": settings.get("tg_auth_verified", False),
        "first_name": settings.get("first_name", ""),
        "last_name": settings.get("last_name", ""),
        "middle_name": settings.get("middle_name", ""),
        "email": settings.get("email", ""),
        "tg_login_code": settings.get("tg_login_code", ""),
        "tg_login_code_expires_at": settings.get("tg_login_code_expires_at", ""),
        "ui_ready": False,
        "all_schedule_groups": normalize_group_names(group_options if group_options else []),
    }

    current_email = (state.get("email") or "").strip().lower()
    if current_email:
        if current_email == DEFAULT_ADMIN_EMAIL:
            state["user_role"] = "Адміністратор"
            save_user_role_to_db("Адміністратор")
            set_user_role_by_email(current_email, "Адміністратор")
        role_from_db = get_user_role_by_email(current_email)
        if role_from_db and role_from_db != state.get("user_role"):
            state["user_role"] = role_from_db
            save_user_role_to_db(role_from_db)

    def persist_bot_settings():
        save_telegram_settings(
            state.get("telegram_enabled", False),
            state.get("telegram_bot_token", ""),
            state.get("telegram_chat_id", ""),
        )

    def apply_bot_updates(updates: dict):
        changed = False
        if "notify_group" in updates:
            ng = (updates.get("notify_group") or "").strip()
            if state.get("notify_group") != ng:
                state["notify_group"] = ng
                changed = True
        if "notifications_enabled" in updates:
            nv = bool(updates.get("notifications_enabled"))
            if state.get("notifications_enabled") != nv:
                state["notifications_enabled"] = nv
                changed = True
        if "notify_minutes" in updates:
            try:
                nm = max(1, min(180, int(updates.get("notify_minutes"))))
            except (TypeError, ValueError):
                nm = state.get("notify_minutes", 10)
            if state.get("notify_minutes") != nm:
                state["notify_minutes"] = nm
                changed = True
        if "telegram_chat_id" in updates:
            chat_id = (updates.get("telegram_chat_id") or "").strip()
            if state.get("telegram_chat_id") != chat_id:
                state["telegram_chat_id"] = chat_id
                changed = True
                persist_bot_settings()
        if "tg_login_code" in updates:
            state["tg_login_code"] = (updates.get("tg_login_code") or "").strip()
        if "tg_login_code_expires_at" in updates:
            state["tg_login_code_expires_at"] = (updates.get("tg_login_code_expires_at") or "").strip()
        if "tg_login_verified" in updates:
            verified = bool(updates.get("tg_login_verified"))
            state["tg_auth_verified"] = verified
            if verified:
                state["auth_required"] = False
            save_telegram_auth_profile(
                auth_required=state.get("auth_required", True),
                verified=state.get("tg_auth_verified", False),
                first_name=state.get("first_name", ""),
                last_name=state.get("last_name", ""),
                middle_name=state.get("middle_name", ""),
                email=state.get("email", ""),
                login_code=state.get("tg_login_code", ""),
                login_code_expires_at=state.get("tg_login_code_expires_at", ""),
            )
            changed = True
        if "set_role_for_email" in updates:
            payload = updates.get("set_role_for_email") or {}
            role = (payload.get("role") or "Студент").strip()
            email = (payload.get("email") or "").strip().lower()
            if email:
                set_user_role_by_email(email, role)
                if email == (state.get("email") or "").strip().lower():
                    state["user_role"] = role
                    save_user_role_to_db(role)
                    changed = True
        if updates.get("parse_groups"):
            trigger_full_groups_parse(show_feedback=False)
        if "user_role" in updates:
            new_role = (updates.get("user_role") or "Студент").strip()
            allowed_roles = {"Студент", "Викладач", "Адміністратор"}
            if new_role not in allowed_roles:
                new_role = "Студент"
            if state.get("user_role") != new_role:
                state["user_role"] = new_role
                save_user_role_to_db(new_role)
                changed = True
        if "announcement_text" in updates:
            text = (updates.get("announcement_text") or "").strip()
            if text:
                save_announcement(text)
                telegram_bridge.notify(f"📢 Повідомлення від бота:\n{text}")
                if state.get("ui_ready"):
                    show_system_notification("Нове повідомлення", text[:200], timeout=8)
                    show_snack(page, f"📢 {text}")
                changed = True
        if "announcement_payload" in updates:
            payload = updates.get("announcement_payload") or {}
            title = (payload.get("title") or "").strip()
            text = (payload.get("text") or "").strip()
            target_group = (payload.get("target_group") or "").strip()
            target_groups = [g.strip() for g in (payload.get("target_groups") or []) if (g or "").strip()]
            author_name = (payload.get("author_name") or "").strip()
            author_role = (payload.get("author_role") or "").strip()
            if title or text:
                save_announcement(
                    text,
                    title=title,
                    target_group=target_group,
                    target_groups=target_groups,
                    author_name=author_name,
                    author_role=author_role,
                )
                preview_parts = []
                if title:
                    preview_parts.append(f"📢 {title}")
                if text:
                    preview_parts.append(text)
                groups_preview = target_groups or ([target_group] if target_group else [])
                if groups_preview:
                    preview_parts.append(f"Групи: {', '.join(groups_preview)}")
                if author_name:
                    preview_parts.append(f"Автор: {author_name}")
                telegram_bridge.notify("\n".join(preview_parts))
                if state.get("ui_ready"):
                    local_msg = "\n".join(preview_parts[:3]) if preview_parts else "Нове повідомлення"
                    show_system_notification("Нове повідомлення", local_msg[:200], timeout=8)
                    show_snack(page, f"📢 {local_msg}")
                changed = True
        if changed:
            save_settings_to_db(
                state["notifications_enabled"],
                state["notify_minutes"],
                state["notify_group"],
                state["theme"],
            )
            try:
                if state.get("auth_required") and not state.get("tg_auth_verified"):
                    page.clean()
                    render_auth_gate()
                else:
                    update_content()
            except Exception:
                pass

    telegram_bridge = TelegramBridge(
        get_state=lambda: state,
        apply_updates=apply_bot_updates,
        get_groups=lambda: list(get_all_available_groups()),
    )

    def current_target_group() -> str:
        return (state.get("notify_group") or state.get("selected_group") or "").strip()

    def get_group_schedule(group_name: str) -> tuple[str, dict]:
        parsed = state.get("parsed_schedule_per_group", {}) or {}
        gname = (group_name or "").strip()
        if not gname or not isinstance(parsed, dict):
            return "", {}
        if gname in parsed:
            return gname, parsed.get(gname, {}) or {}
        target_codes = set(normalize_group_names([gname]))
        if not target_codes:
            return "", {}
        for key, val in parsed.items():
            key_codes = set(normalize_group_names([key]))
            if target_codes.intersection(key_codes):
                return key, val or {}
        return "", {}

    def get_all_available_groups() -> list[str]:
        groups = set(normalize_group_names(state.get("group_options") or []))
        groups.update(normalize_group_names(state.get("all_schedule_groups") or []))
        parsed = state.get("parsed_schedule_per_group") or {}
        if isinstance(parsed, dict):
            for g in normalize_group_names(list(parsed.keys())):
                groups.add(g)
        cleaned: set[str] = set()
        for g in groups:
            for code in _canonical_group_codes(g):
                cleaned.add(code)
        return sorted(cleaned)

    def _extract_groups_from_file_name(name: str) -> list[str]:
        base = (name or "").strip()
        if not base:
            return []
        base = re.sub(r"\.(pdf|docx?)$", "", base, flags=re.IGNORECASE).strip()
        chunks = [c.strip() for c in re.split(r"[,\n;/]+", base) if c.strip()]
        out = []
        for ch in chunks:
            cleaned = re.sub(r"\([^)]*\)", "", ch).strip()
            out.extend(_canonical_group_codes(cleaned))
        return normalize_group_names(out)

    def _extract_groups_from_file_content_bytes(file_name: str, file_url: str, data: bytes) -> list[str]:
        try:
            lname = (file_name or file_url or "").lower()
            text_blob = ""
            if lname.endswith(".docx"):
                doc = Document(io.BytesIO(data))
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = (cell.text or "").strip()
                            if cell_text:
                                text_blob += "\n" + cell_text
                for p in doc.paragraphs:
                    pt = (p.text or "").strip()
                    if pt:
                        text_blob += "\n" + pt
            elif lname.endswith(".pdf"):
                pdf = fitz.open(stream=data, filetype="pdf")
                max_pages = min(4, len(pdf))
                for i in range(max_pages):
                    try:
                        text_blob += "\n" + (pdf[i].get_text("text") or "")
                    except Exception:
                        continue
                pdf.close()
            else:
                return []
            return normalize_group_names(_canonical_group_codes(text_blob))
        except Exception:
            return []

    def collect_all_groups_from_schedule_tree(root_url: str) -> list[str]:
        parser = get_parser()
        if not parser or not root_url:
            return []
        visited: set[str] = set()
        found: set[str] = set()

        def walk(url: str, depth: int = 0):
            if depth > 16 or not url or url in visited:
                return
            visited.add(url)
            try:
                items = parser.parse_files(url)
            except Exception:
                return
            for it in items or []:
                if it.is_folder:
                    walk(it.url, depth + 1)
                else:
                    for g in _extract_groups_from_file_name(getattr(it, "name", "")):
                        found.add(g)
                    try:
                        lname = (getattr(it, "name", "") or "").lower()
                        if lname.endswith(".docx") or lname.endswith(".pdf"):
                            response = parser.session.get(it.url, timeout=REQUEST_TIMEOUT)
                            response.raise_for_status()
                            for g in _extract_groups_from_file_content_bytes(getattr(it, "name", ""), it.url, response.content):
                                found.add(g)
                    except Exception:
                        continue

        walk(root_url, 0)
        return sorted(found)

    def trigger_full_groups_parse(show_feedback: bool = True):
        async def _task():
            groups = await asyncio.to_thread(collect_all_groups_from_schedule_tree, state.get("schedule_root_url", ""))
            if groups:
                merged = sorted(set(normalize_group_names((state.get("all_schedule_groups") or []) + groups)))
                state["all_schedule_groups"] = merged
                if show_feedback:
                    show_snack(page, f"Парсинг завершено. Знайдено груп: {len(merged)}")
            elif show_feedback:
                show_snack(page, "Не вдалося знайти групи. Перевірте доступ до розкладу.")

        page.run_task(_task)

    def current_user_fio() -> str:
        return " ".join(
            [
                (state.get("last_name") or "").strip(),
                (state.get("first_name") or "").strip(),
                (state.get("middle_name") or "").strip(),
            ]
        ).strip()

    def find_upcoming_entries() -> list[dict]:
        messages: list[str] = []

        try:

            current_date = datetime.now().date()
            if "last_notification_date" not in state or state.get("last_notification_date") != current_date:
                state["notified_entries"] = get_shown_notifications_for_today()
                state["last_notification_date"] = current_date

            if not state.get("notifications_enabled", True):
                print("[NOTIFICATION] Сповіщення вимкнені")
                return messages

            sg = state.get("notify_group") or state.get("selected_group")
            if not sg:
                print("[NOTIFICATION] Група не вибрана")
                return messages

            all_groups_schedule = state.get("parsed_schedule_per_group", {})
            if not all_groups_schedule:
                print("[NOTIFICATION] Розклад не завантажено")
                return messages

            _, group_schedule = get_group_schedule(sg)
            if not group_schedule:
                print(f"[NOTIFICATION] Розклад для групи '{sg}' не знайдено")
                return messages

            notify_time = int(state.get("notify_minutes", 120)) or 1
            now = datetime.now()
            today_date = now.date()

            print(f"[NOTIFICATION] Перевірка: група={sg}, час={notify_time} хв, сьогодні={today_date.strftime('%d.%m.%Y')}")

            total_entries_checked = 0
            today_entries_found = 0
            for week_name, entries in group_schedule.items():
                if not entries:
                    continue

                for entry in entries:
                    total_entries_checked += 1
                    if not isinstance(entry, dict):
                        continue

                    date_str = entry.get("date", "").strip()
                    time_str = entry.get("time", "").strip()
                    subject = entry.get("subject", "").strip()
                    teacher = entry.get("teacher", "").strip()

                    if not date_str or not time_str or not subject:
                        continue

                    class_date = None
                    try:
                        parts = date_str.split(".")
                        if len(parts) == 3:
                            day, month, year_str = parts

                            if len(year_str) == 2:
                                year_2digit = int(year_str)

                                current_year_2digit = today_date.year % 100
                                if year_2digit <= current_year_2digit:

                                    if current_year_2digit - year_2digit > 1:
                                        full_year = (today_date.year // 100) * 100 + year_2digit
                                    else:

                                        full_year_prev = (today_date.year // 100 - 1) * 100 + year_2digit
                                        full_year_curr = (today_date.year // 100) * 100 + year_2digit

                                        if abs(full_year_prev - today_date.year) < abs(full_year_curr - today_date.year):
                                            full_year = full_year_prev
                                        else:
                                            full_year = full_year_curr
                                else:

                                    full_year = (today_date.year // 100) * 100 + year_2digit
                                class_date = date(int(full_year), int(month), int(day))
                            else:

                                class_date = datetime.strptime(date_str, "%d.%m.%Y").date()
                    except Exception:
                        try:

                            if len(date_str.split(".")[-1]) == 2:
                                class_date = datetime.strptime(date_str, "%d.%m.%y").date()

                                if class_date.year < today_date.year - 1:
                                    class_date = class_date.replace(year=class_date.year + 100)
                            else:
                                class_date = datetime.strptime(date_str, "%d.%m.%Y").date()
                        except Exception:
                            continue

                    if not class_date:
                        continue

                    if class_date != today_date:
                        continue

                    today_entries_found += 1

                    try:

                        if "-" in time_str:
                            time_str = time_str.split("-")[0].strip()

                        if ":" not in time_str:
                            continue
                        time_parts = time_str.split(":")
                        if len(time_parts) != 2:
                            continue
                        hour = int(time_parts[0].strip())
                        minute = int(time_parts[1].strip())

                        if hour < 0 or hour > 23 or minute < 0 or minute > 59:
                            continue
                    except (ValueError, IndexError):
                        continue

                    class_dt = datetime(class_date.year, class_date.month, class_date.day, hour, minute)
                    delta = class_dt - now

                    if delta.total_seconds() <= 0:
                        continue

                    total_seconds = delta.total_seconds()
                    minutes_until = int(math.ceil(total_seconds / 60))
                    seconds_remainder = int(total_seconds % 60)

                    if minutes_until <= notify_time + 5:
                        print(f"[NOTIFICATION] Пара на сьогодні: {time_str} - {subject}, залишилось {minutes_until} хв (потрібно {notify_time} хв)")

                    key = (date_str, time_str, subject)
                    already_shown = key in state["notified_entries"]

                    total_seconds_until = int(total_seconds)
                    notify_seconds = notify_time * 60

                    time_in_range = (notify_seconds <= total_seconds_until <= notify_seconds + 5) and total_seconds_until > 0

                    if minutes_until <= notify_time + 3 and minutes_until > 0:
                        print(f"[NOTIFICATION] Деталі: {time_str} - {subject}")
                        print(f"  - Залишилось: {minutes_until} хв {seconds_remainder} сек ({total_seconds_until} сек, потрібно {notify_seconds} сек)")
                        print(f"  - Час у діапазоні: {time_in_range} (допуск +5 сек вперед)")
                        print(f"  - Вже показано: {already_shown}")
                        if already_shown:
                            print(f"  - ⚠️ Уведомлення вже було показано раніше!")

                    if time_in_range and not already_shown:
                        print(f"[NOTIFICATION] ✅ Знайдено пару для сповіщення: {time_str} - {subject}, залишилось {minutes_until} хв {seconds_remainder} сек")

                        if minutes_until >= 60:
                            hours_left = minutes_until // 60
                            minutes_remainder = minutes_until % 60
                            time_str_notify = f"{hours_left} год {minutes_remainder} хв" if minutes_remainder > 0 else f"{hours_left} год"
                        else:
                            time_str_notify = f"{minutes_until} хв"

                        meet_link = None
                        subject_lower = subject.lower()
                        for key, link in MEET_LINKS.items():
                            if key.lower() in subject_lower:
                                meet_link = link
                                break

                        teacher_text = f", викладач: {teacher}" if teacher else ""
                        msg = f"Через {time_str_notify} починається пара ({time_str}) у групі {sg}: {subject}{teacher_text}"

                        notification_data = {
                            "message": msg,
                            "subject": subject,
                            "teacher": teacher,
                            "time": time_str,
                            "group": sg,
                            "link": meet_link
                        }
                        messages.append(notification_data)

                        try:
                            mark_notification_shown(date_str, time_str, subject)
                            state["notified_entries"].add(key)
                            save_announcement(msg)
                            telegram_bridge.notify(f"🔔 {msg}")
                            print(f"[NOTIFICATION] Показано уведомлення: {msg}")
                        except Exception as e:
                            print(f"[ERROR] Помилка збереження уведомлення: {e}")

            if total_entries_checked > 0:
                print(f"[NOTIFICATION] Перевірено {total_entries_checked} пар, знайдено {today_entries_found} на сьогодні, знайдено {len(messages)} уведомлень")
            elif total_entries_checked == 0:
                print(f"[NOTIFICATION] У розкладі групи '{sg}' немає пар")

        except Exception as ex:
            print(f"[ERROR] Помилка в find_upcoming_entries: {ex}")
            import traceback
            traceback.print_exc()

        return messages

    def check_notifications(e):
        notifications = find_upcoming_entries()
        for notif_data in notifications:
            show_class_notification(page, notif_data, colors)

    def check_notifications_dialog(e):
        try:
            group_filter = current_target_group()
            all_announcements = get_all_announcements(target_group=group_filter, include_meta=True)
            unread_announcements = get_unread_announcements(target_group=group_filter, include_meta=True)
            unread_ids = {aid for aid, _, _, _ in unread_announcements}

            if not all_announcements:
                content = ft.Text("Немає збережених повідомлень.", size=16, color=colors["TEXT_MUTED"])
            else:
                announcement_items = []
                for aid, msg, created_at, _meta in all_announcements:
                    msg_clean = (msg or "").strip()
                    if msg_clean:
                        is_unread = aid in unread_ids
                        bg_color = colors["ACCENT"] if is_unread else colors["BG_CARD"]
                        border_color = colors["ACCENT"] if is_unread else colors["BORDER_COLOR"]

                        date_str = created_at if created_at else "Дата невідома"
                        unread_badge = ft.Container(
                            content=ft.Text("НОВЕ", size=10, weight="bold", color=ft.colors.WHITE),
                            bgcolor=colors["ACCENT"],
                            padding=ft.Padding(6, 2, 6, 2),
                            border_radius=4,
                        ) if is_unread else None

                        header_row = [ft.Text(f"📢 {date_str}", size=12, color=colors["TEXT_MUTED"] if not is_unread else ft.colors.WHITE, weight="bold")]
                        if unread_badge:
                            header_row.append(unread_badge)

                        item_container = ft.Container(
                            content=ft.Column(
                                [
                                    ft.Row(header_row, spacing=10),
                                    ft.Text(msg_clean, size=14, color=ft.colors.WHITE if is_unread else colors["TEXT_COLOR"]),
                                    ft.Divider(thickness=1, color=border_color, height=10),
                                ],
                                spacing=5,
                            ),
                            padding=ft.Padding(10, 8, 10, 8),
                            bgcolor=bg_color,
                            border=_border_all(2, border_color) if is_unread else None,
                            border_radius=8 if is_unread else 0,
                        )

                        if is_unread:
                            def make_mark_read_handler(announcement_id):
                                def mark_read_click(ev):
                                    mark_announcements_seen(announcement_id)
                                    check_notifications_dialog(ev)
                                    check_button_ref = state.get("check_button_ref")
                                    if check_button_ref and check_button_ref.current:
                                        unread_count = len(get_unread_announcements(target_group=current_target_group()))
                                        button_text = "Перевірити повідомлення"
                                        if unread_count > 0:
                                            button_text += f" ({unread_count} нових)"
                                        check_button_ref.current.text = button_text
                                        check_button_ref.current.bgcolor = colors["ACCENT"] if unread_count > 0 else colors["BORDER_COLOR"]
                                        check_button_ref.current.update()
                                return mark_read_click
                            item_container.on_click = make_mark_read_handler(aid)

                        announcement_items.append(item_container)

                content = ft.Column(announcement_items, spacing=5, scroll=ft.ScrollMode.AUTO)
        except Exception as ex:
            content = ft.Text(f"Помилка при завантаженні повідомлень: {str(ex)}", size=14, color=ft.colors.RED)

        dialog = ft.AlertDialog(
            title=ft.Text(f"Повідомлення ({len(unread_ids)} нових)" if unread_ids else "Всі повідомлення", size=18, weight="bold"),
            content=ft.Container(
                content=content,
                padding=20,
                width=500,
                height=400,
            ),
            actions=[],
            actions_alignment=ft.MainAxisAlignment.CENTER,
        )

        if unread_ids:
            max_unread_id = max(unread_ids)
            def mark_all_read(e):
                mark_announcements_seen(max_unread_id)
                dialog.open = False
                page.update()
                check_button_ref = state.get("check_button_ref")
                if check_button_ref and check_button_ref.current:
                    check_button_ref.current.text = "Перевірити повідомлення"
                    check_button_ref.current.bgcolor = colors["BORDER_COLOR"]
                    check_button_ref.current.update()

            dialog.actions = [
                ft.Button(
                    "Позначити всі як прочитані",
                    on_click=mark_all_read,
                    bgcolor=colors["ACCENT"],
                    style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
                ),
                ft.Button(
                    "Закрити",
                    on_click=lambda ev: setattr(dialog, 'open', False) or page.update(),
                    bgcolor=colors["BORDER_COLOR"],
                    style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
                ),
            ]
        else:
            dialog.actions = [
                ft.Button(
                    "OK",
                    on_click=lambda ev: setattr(dialog, 'open', False) or page.update(),
                    bgcolor=colors["ACCENT"],
                    style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
                )
            ]

        page.dialog = dialog
        dialog.open = True
        page.update()

    def close_dialog(e, dialog: ft.AlertDialog):
        dialog.open = False
        page.update()

    def messages_page():
        group_filter = current_target_group()
        all_announcements = get_all_announcements(target_group=group_filter, include_meta=True)
        unread_announcements = get_unread_announcements(target_group=group_filter, include_meta=True)
        unread_ids = {aid for aid, _, _, _ in unread_announcements}

        selected_ids = set()
        delete_button_ref = ft.Ref[ft.Button]()
        select_all_button_ref = ft.Ref[ft.Button]()
        messages_list_ref = ft.Ref[ft.Column]()
        checkboxes_refs = {}

        def update_delete_button():
            try:
                count = len(selected_ids)
                if delete_button_ref.current and delete_button_ref.current.page:
                    delete_button_ref.current.text = f"Видалити вибрані ({count})" if count > 0 else "Видалити вибрані"
                    delete_button_ref.current.disabled = count == 0
                    delete_button_ref.current.update()
            except:
                pass

        def update_select_all_button():
            try:
                total_ids = {aid for aid, _, _, _ in all_announcements}
                all_selected = len(selected_ids) == len(total_ids) and len(total_ids) > 0
                if select_all_button_ref.current and select_all_button_ref.current.page:
                    select_all_button_ref.current.text = "Скасувати вибір" if all_selected else "Вибрати все"
                    select_all_button_ref.current.update()
            except:
                pass

        def toggle_select_all(e):
            try:
                total_ids = {aid for aid, _, _, _ in all_announcements}
                all_selected = len(selected_ids) == len(total_ids) and len(total_ids) > 0

                if all_selected:
                    selected_ids.clear()
                    for checkbox_ref in checkboxes_refs.values():
                        if checkbox_ref.current and checkbox_ref.current.page:
                            checkbox_ref.current.value = False
                            checkbox_ref.current.update()
                else:
                    selected_ids.update(total_ids)
                    for checkbox_ref in checkboxes_refs.values():
                        if checkbox_ref.current and checkbox_ref.current.page:
                            checkbox_ref.current.value = True
                            checkbox_ref.current.update()

                update_delete_button()
                update_select_all_button()
            except Exception as ex:
                print(f"Error in toggle_select_all: {ex}")

        def refresh_messages():
            selected_ids.clear()
            checkboxes_refs.clear()
            update_delete_button()
            update_content()

        if not all_announcements:
            content_list = [
                ft.Container(
                    content=ft.Text("Немає збережених повідомлень.", size=16, color=colors["TEXT_MUTED"]),
                    padding=40,
                    alignment=ft.Alignment(0, 0),
                )
            ]
        else:
            content_list = []

            header_buttons = []

            header_buttons.append(
                ft.Button(
                    "Вибрати все",
                    ref=select_all_button_ref,
                    on_click=toggle_select_all,
                    bgcolor=colors["BORDER_COLOR"],
                    width=180,
                    style=ft.ButtonStyle(
                        padding=_padding_symmetric(horizontal=24, vertical=14),
                        shape=ft.RoundedRectangleBorder(radius=10)
                    ),
                )
            )

            if unread_ids:
                def mark_all_read(e):
                    max_unread_id = max(unread_ids)
                    mark_announcements_seen(max_unread_id)
                    refresh_messages()

                header_buttons.append(
                    ft.Button(
                        "Позначити всі як прочитані",
                        on_click=mark_all_read,
                        bgcolor=colors["ACCENT"],
                        width=220,
                        style=ft.ButtonStyle(
                            padding=_padding_symmetric(horizontal=24, vertical=14),
                            shape=ft.RoundedRectangleBorder(radius=10)
                        ),
                    )
                )

            def delete_selected(e):
                if selected_ids:
                    delete_announcements(list(selected_ids))
                    refresh_messages()

            header_buttons.append(
                ft.Button(
                    "Видалити вибрані",
                    ref=delete_button_ref,
                    on_click=delete_selected,
                    bgcolor=ft.colors.RED_600,
                    disabled=True,
                    width=180,
                    style=ft.ButtonStyle(
                        padding=_padding_symmetric(horizontal=24, vertical=14),
                        shape=ft.RoundedRectangleBorder(radius=10)
                    ),
                )
            )

            content_list.append(
                ft.Container(
                    content=ft.Row(
                        [
                            ft.Text(
                                f"Повідомлення ({len(unread_ids)} нових)" if unread_ids else "Всі повідомлення",
                                size=20,
                                weight="bold",
                                color=colors["TEXT_COLOR"],
                                expand=True,
                            ),
                            ft.Column(
                                header_buttons,
                                spacing=10,
                                horizontal_alignment=ft.CrossAxisAlignment.END,
                            ),
                        ],
                        alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                    ),
                    padding=ft.Padding(0, 0, 0, 20),
                )
            )

            announcement_items = []
            for aid, msg, created_at, meta in all_announcements:
                msg_clean = (msg or "").strip()
                if msg_clean:
                    is_unread = aid in unread_ids
                    bg_color = colors["ACCENT"] if is_unread else colors["BG_CARD"]
                    border_color = colors["ACCENT"] if is_unread else colors["BORDER_COLOR"]

                    date_str = created_at if created_at else "Дата невідома"
                    unread_badge = ft.Container(
                        content=ft.Text("НОВЕ", size=10, weight="bold", color=ft.colors.WHITE),
                        bgcolor=colors["ACCENT"],
                        padding=ft.Padding(6, 2, 6, 2),
                        border_radius=4,
                    ) if is_unread else None

                    checkbox_ref = ft.Ref[ft.Checkbox]()
                    checkboxes_refs[aid] = checkbox_ref

                    def make_checkbox_handler(announcement_id):
                        def on_change(e):
                            try:
                                if e.control.value:
                                    selected_ids.add(announcement_id)
                                else:
                                    selected_ids.discard(announcement_id)
                                update_delete_button()
                                update_select_all_button()
                            except:
                                pass
                        return on_change

                    checkbox = ft.Checkbox(
                        ref=checkbox_ref,
                        value=aid in selected_ids,
                        on_change=make_checkbox_handler(aid),
                    )

                    header_row = [
                        checkbox,
                        ft.Text(f"📢 {date_str}", size=12, color=colors["TEXT_MUTED"] if not is_unread else ft.colors.WHITE, weight="bold"),
                    ]
                    if unread_badge:
                        header_row.append(unread_badge)

                    def make_mark_read_handler(announcement_id):
                        def mark_read_click(ev):
                            mark_announcements_seen(announcement_id)
                            refresh_messages()
                        return mark_read_click

                    details_controls = [ft.Text(msg_clean, size=14, color=ft.colors.WHITE if is_unread else colors["TEXT_COLOR"])]
                    attachment_name = (meta.get("attachment_name") or "").strip()
                    attachment_path = (meta.get("attachment_path") or "").strip()
                    if attachment_name:
                        def open_attachment(path: str):
                            def _handler(ev):
                                try:
                                    if path and os.path.exists(path):
                                        page.launch_url("file://" + path, web_popup_window=False)
                                    else:
                                        show_snack(page, "Файл не знайдено. Можливо, його переміщено.")
                                except Exception:
                                    show_snack(page, "Не вдалося відкрити файл.")
                            return _handler
                        details_controls.append(
                            ft.TextButton(
                                f"Відкрити файл: {attachment_name}",
                                on_click=open_attachment(attachment_path),
                                style=ft.ButtonStyle(color=colors["ACCENT_TEXT"] if is_unread else colors["ACCENT"]),
                            )
                        )
                    item_container = ft.Container(
                        content=ft.Column(
                            [
                                ft.Row(header_row, spacing=10),
                                *details_controls,
                                ft.Divider(thickness=1, color=border_color, height=10),
                            ],
                            spacing=5,
                        ),
                        padding=ft.Padding(10, 8, 10, 8),
                        bgcolor=bg_color,
                        border=_border_all(2, border_color) if is_unread else None,
                        border_radius=8 if is_unread else 0,
                        margin=ft.Margin(0, 0, 0, 10),
                    )

                    def make_click_handler(announcement_id):
                        def on_click(ev):
                            if hasattr(ev, 'control') and isinstance(ev.control, ft.Checkbox):
                                return
                            if announcement_id in unread_ids:
                                mark_announcements_seen(announcement_id)
                                refresh_messages()
                        return on_click
                    item_container.on_click = make_click_handler(aid)

                    announcement_items.append(item_container)

            content_list.extend(announcement_items)

        messages_column = ft.Column(
            content_list,
            spacing=0,
            scroll=ft.ScrollMode.AUTO,
            expand=True,
            ref=messages_list_ref,
        )

        return ft.Container(
            content=messages_column,
            padding=24,
            expand=True,
        )

    news_list_ref = ft.Ref[ft.ListView]()

    def news_page():
        if "news_items" not in state:
            state["news_items"] = []
        if "news_next_page" not in state:
            state["news_next_page"] = 1
        if "news_loading" not in state:
            state["news_loading"] = False
        if "news_no_more" not in state:
            state["news_no_more"] = False
        if "news_refreshing" not in state:
            state["news_refreshing"] = False
        if "news_last_scroll_pixels" not in state:
            state["news_last_scroll_pixels"] = -1
        if not state["news_items"] and not state["news_loading"] and not state.get("news_refreshing"):
            state["news_loading"] = True
            try:
                state["news_items"] = parse_istu_news(1)
                state["news_next_page"] = 2
            except Exception:
                state["news_items"] = []
                state["news_next_page"] = 1
            state["news_loading"] = False
        items = state.get("news_items") or []

        def open_news_url(url: str):
            def handler(e):
                try:
                    page.launch_url(url, web_popup_window=False)
                except Exception as ex:
                    try:
                        webbrowser.open(url)
                    except Exception:
                        show_snack(page, f"Не вдалося відкрити посилання: {ex}")
            return handler

        def load_more_news():
            if state.get("selected_index", 0) != 0:
                return
            if state.get("news_loading") or state.get("news_no_more"):
                return
            saved_scroll = state.get("news_last_scroll_pixels", 0)
            state["news_loading"] = True
            try:
                next_p = state.get("news_next_page", 1)
                more = []
                for p in (next_p, next_p + 1):
                    chunk = parse_istu_news(p)
                    if chunk and not (len(chunk) == 1 and chunk[0].get("title") == "Помилка завантаження"):
                        more.extend(chunk)
                    else:
                        break
                if more:
                    state["news_items"] = (state.get("news_items") or []) + more
                    state["news_next_page"] = next_p + 2
                else:
                    state["news_no_more"] = True
            except Exception:
                pass
            state["news_loading"] = False
            if state.get("selected_index", 0) == 0:
                update_content()
                if news_list_ref.current and saved_scroll > 0:
                    try:
                        news_list_ref.current.scroll_to(offset=saved_scroll, duration=0)
                    except Exception:
                        pass
                    page.update()

        def refresh_feed():
            if state.get("news_refreshing"):
                return
            state["news_refreshing"] = True
            if state.get("selected_index", 0) == 0:
                update_content()
            try:
                fresh = parse_istu_news(1)
                if fresh and not (len(fresh) == 1 and fresh[0].get("title") == "Помилка завантаження"):
                    state["news_items"] = fresh
                    state["news_next_page"] = 2
                    state["news_no_more"] = False
            except Exception:
                pass
            state["news_refreshing"] = False
            if state.get("selected_index", 0) == 0:
                update_content()

        def on_scroll(e):
            if state.get("selected_index", 0) != 0:
                return
            try:
                pixels = getattr(e, "pixels", None)
                max_extent = getattr(e, "max_scroll_extent", None)
                if pixels is None and getattr(e, "data", None):
                    try:
                        d = json.loads(e.data) if isinstance(e.data, str) else e.data
                        pixels = d.get("pixels", 0)
                        max_extent = d.get("max_scroll_extent")
                    except Exception:
                        pixels, max_extent = 0, None
                pixels = pixels if pixels is not None else 0
                last = state.get("news_last_scroll_pixels", -1)
                state["news_last_scroll_pixels"] = pixels
                if max_extent is not None and max_extent > 100 and pixels >= max_extent - 400:
                    load_more_news()
                elif pixels <= 80 and last > 150:
                    refresh_feed()
            except Exception:
                pass

        header = [
            ft.Row(
                [
                    ft.Text("Новини МНТУ", size=22, weight="bold", color=colors["TEXT_COLOR"]),
                    ft.Button(
                        "Оновити стрічку",
                        icon=ft.icons.REFRESH,
                        on_click=lambda e: refresh_feed(),
                        bgcolor=colors["ACCENT"],
                        style=ft.ButtonStyle(color=colors["ACCENT_TEXT"], padding=_padding_symmetric(horizontal=14, vertical=10), shape=ft.RoundedRectangleBorder(radius=10)),
                    ),
                ],
                alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                wrap=True,
            ),
            ft.Container(height=8),
            ft.Text("Вгору — оновлення. Вниз — старіші новини. Або натисніть кнопку вище.", size=12, color=colors["TEXT_MUTED"]),
            ft.Container(height=16),
        ]
        list_controls = list(header)
        if not items:
            list_controls.append(
                ft.Container(
                    content=ft.Text("Не вдалося завантажити новини з istu.edu.ua/novyny/", size=16, color=colors["TEXT_MUTED"]),
                    alignment=ft.Alignment(0, 0),
                    padding=40,
                )
            )
        else:
            for item in items:
                title = (item.get("title") or "Новина").strip()
                url = item.get("url") or NEWS_URL
                date_str = (item.get("date") or "").strip()
                excerpt = (item.get("excerpt") or "").strip()
                img_url = (item.get("image") or "").strip()
                is_pinned = bool(item.get("is_pinned"))
                card_content = []
                if img_url:
                    card_content.append(
                        ft.Container(
                            content=ft.Image(
                                src=img_url,
                                width=400,
                                height=220,
                                fit=ft.ImageFit.COVER,
                                border_radius=ft.BorderRadius(12, 12, 0, 0),
                            ),
                            clip_behavior=ft.ClipBehavior.HARD_EDGE,
                        )
                    )
                    card_content.append(ft.Container(height=12))
                if is_pinned:
                    card_content.append(
                        ft.Container(
                            content=ft.Row(
                                [
                                    ft.Icon(ft.icons.PUSH_PIN, size=14, color=colors["ACCENT_TEXT"]),
                                    ft.Text("Закріплено", size=12, weight="bold", color=colors["ACCENT_TEXT"]),
                                ],
                                spacing=6,
                                tight=True,
                            ),
                            padding=_padding_symmetric(horizontal=10, vertical=6),
                            border_radius=999,
                            bgcolor=colors["ACCENT"],
                        )
                    )
                    card_content.append(ft.Container(height=8))
                card_content.extend([
                    ft.Text(title, size=16, weight="bold", color=colors["TEXT_COLOR"]),
                    ft.Container(height=4),
                    ft.Row(
                        [
                            ft.Text(date_str, size=12, color=colors["TEXT_MUTED"]) if date_str else ft.Container(),
                            ft.TextButton("Читати далі →", on_click=open_news_url(url), style=ft.ButtonStyle(color=ft.colors.BLUE)),
                        ],
                        alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                    ),
                    ft.Container(height=6),
                    ft.Text(excerpt, size=14, color=colors["TEXT_MUTED"], max_lines=8, overflow=ft.TextOverflow.ELLIPSIS) if excerpt else ft.Container(),
                ])
                card = ft.Container(
                    content=ft.Column(
                        card_content,
                        spacing=0,
                        alignment=ft.MainAxisAlignment.START,
                        horizontal_alignment=ft.CrossAxisAlignment.START,
                    ),
                    padding=16,
                    border_radius=12,
                    bgcolor=colors["BG_CARD"],
                    border=_border_all(1, colors["BORDER_COLOR"]),
                )
                list_controls.append(card)
                list_controls.append(ft.Container(height=12))
            if state.get("news_no_more"):
                load_more_hint = ft.Container(
                    height=50,
                    content=ft.Text("Усі новини завантажено", size=12, color=colors["TEXT_MUTED"]),
                    alignment=ft.Alignment(0, 0),
                )
            else:
                load_more_hint = ft.Container(
                    height=70,
                    content=ft.Column(
                        [
                            ft.ProgressRing(width=32, height=32, stroke_width=2) if state.get("news_loading") else ft.Button(
                                "Завантажити ще",
                                icon=ft.icons.KEYBOARD_ARROW_DOWN,
                                on_click=lambda e: load_more_news(),
                                bgcolor=colors["ACCENT"],
                                style=ft.ButtonStyle(color=colors["ACCENT_TEXT"], padding=_padding_symmetric(horizontal=20, vertical=12), shape=ft.RoundedRectangleBorder(radius=10)),
                            ),
                            ft.Container(height=4),
                            ft.Text("Прокрутіть вниз — старіші новини", size=12, color=colors["TEXT_MUTED"]),
                        ],
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        spacing=0,
                    ),
                    alignment=ft.Alignment(0, 0),
                )
            list_controls.append(load_more_hint)

        list_view = ft.ListView(
            list_controls,
            expand=True,
            spacing=0,
            padding=24,
            on_scroll=on_scroll,
            auto_scroll=False,
            ref=news_list_ref,
        )
        return ft.Container(content=list_view, expand=True)

    def schedule_page():
        if (
            state.get("week_options")
            and state.get("group_options")
            and schedule_has_entries(state.get("parsed_schedule_per_group") or {})
        ):
            all_weeks = state["week_options"]
            sel_week = state["selected_week"]
            all_groups = state["group_options"]
            sel_group = state["selected_group"]

            content_list: list[ft.Control] = [
                ft.Row(
                    [
                        ft.IconButton(
                            icon=ft.icons.ARROW_BACK,
                            icon_color=ft.colors.WHITE,
                            on_click=close_group_view,
                            tooltip="Назад до файлів",
                        ),
                        ft.Text("Розклад (виберіть тиждень і групу)", size=20, weight="bold"),
                    ],
                    alignment=ft.MainAxisAlignment.START,
                ),
                ft.Divider(thickness=1, color=colors["BORDER_COLOR"]),

                ft.Text(
                    "Оберіть тиждень:",
                    size=18,
                    weight="bold",
                    text_align=ft.TextAlign.LEFT,
                ),
                ft.Dropdown(
                    width=200,
                    value=sel_week or None,
                    options=[ft.dropdown.Option(w) for w in all_weeks],
                    hint_text="Тиждень",
                    **_dropdown_event_kw(lambda e: select_week(e)),
                ),
                ft.Divider(thickness=1, color=colors["BORDER_COLOR"]),

                ft.Text(
                    "Оберіть групу:",
                    size=18,
                    weight="bold",
                    text_align=ft.TextAlign.LEFT,
                ),
                ft.Dropdown(
                    width=300,
                    value=sel_group or None,
                    options=[ft.dropdown.Option(g) for g in all_groups],
                    hint_text="Група",
                    **_dropdown_event_kw(lambda e: select_group(e)),
                ),
                ft.Divider(thickness=1, color=colors["BORDER_COLOR"]),
            ]

            if not sel_week or not sel_group:
                content_list.append(
                    ft.Text(
                        "Оберіть спочатку тиждень та групу.",
                        size=16,
                        color=ft.colors.RED,
                        text_align=ft.TextAlign.CENTER,
                    )
                )
            else:
                week_key = sel_week
                today = date.today()
                _, sel_group_schedule = get_group_schedule(sel_group)
                entries = list((sel_group_schedule or {}).get(week_key, []))
                real_entries = [e for e in entries if not _is_synthetic_no_classes_entry(e)]
                mon_cur = today - timedelta(days=today.weekday())
                if week_key == "Цей тиждень":
                    week_human = format_week_range_human(mon_cur, mon_cur + timedelta(days=6))
                elif week_key == "Наступний тиждень":
                    mon_n = mon_cur + timedelta(days=7)
                    week_human = format_week_range_human(mon_n, mon_n + timedelta(days=6))
                else:
                    week_human = str(week_key)

                dates_for_week = sorted(
                    {e["date"] for e in real_entries if (e.get("date") or "").strip()},
                    key=lambda s: parse_schedule_date_str(s) or date.min,
                )
                content_list.append(
                    ft.Text(
                        f"Календарний тиждень: {week_human} · сьогодні {format_schedule_date_short(today)}",
                        size=13,
                        color=colors["TEXT_MUTED"],
                    )
                )
                content_list.append(ft.Container(height=6))

                if dates_for_week:
                    content_list.append(
                        ft.Text(
                            "Розклад за датами",
                            size=18,
                            weight="bold",
                            color=ft.colors.WHITE,
                        )
                    )
                    content_list.append(ft.Divider(thickness=1, color=colors["BORDER_COLOR"]))

                if not real_entries:
                    if entries:
                        placeholder = (entries[0].get("subject") or "").strip() or "Інформації немає"
                        content_list.append(
                            ft.Container(
                                padding=12,
                                content=ft.Text(placeholder, size=16, color=colors["TEXT_MUTED"]),
                            )
                        )
                    else:
                        content_list.append(
                            ft.Container(
                                padding=8,
                                content=ft.Text(
                                    f"На цей календарний тиждень ({week_human}) у завантаженому файлі ще немає розкладу: "
                                    f"за датами цього тижня пар не знайдено (таблицю не оновили або у документі немає цих дат). "
                                    f"Сьогодні {format_schedule_date_short(today)}. Завантажте актуальний файл з lib.istu.edu.ua.",
                                    size=15,
                                    color=ft.colors.AMBER_200,
                                ),
                            )
                        )
                else:

                    days_schedule = {}
                    for day_date in dates_for_week:
                        day_entries = [e for e in real_entries if e["date"] == day_date]
                        if day_entries:

                            time_groups = {}
                            for entry in day_entries:
                                time_key = entry['time']
                                if time_key not in time_groups:
                                    time_groups[time_key] = []

                                subject_info = {
                                    "subject": entry.get('subject', ''),
                                    "teacher": entry.get('teacher', '')
                                }
                                time_groups[time_key].append(subject_info)
                            days_schedule[day_date] = sorted(time_groups.items(), key=lambda x: x[0])

                    sorted_days = [(day_date, days_schedule[day_date]) for day_date in dates_for_week if day_date in days_schedule]

                    for day_idx, (day_date, time_entries) in enumerate(sorted_days):

                        day_title = day_date

                        parsed_day_date = extract_date_from_text(day_date, today)
                        is_today = False
                        has_classes = len(time_entries) > 0

                        if parsed_day_date:
                            is_today = parsed_day_date == today

                        info_parts = []
                        if is_today:
                            info_parts.append("сьогодні")
                        if has_classes:
                            num_classes = len(time_entries)
                            if num_classes == 1:
                                info_parts.append("1 пара")
                            else:
                                info_parts.append(f"{num_classes} пари")
                        else:
                            info_parts.append("пар немає")

                        if info_parts:
                            day_title += f" ({', '.join(info_parts)})"

                        content_list.append(
                            ft.Container(
                                content=ft.Text(
                                    day_title,
                                    size=18,
                                    weight="bold",
                                    color=ft.colors.WHITE
                                ),
                                padding=ft.Padding(0, 15, 0, 10),
                                bgcolor=colors["BG_DARK"],
                            )
                        )

                        content_list.append(
                            ft.Container(
                                content=ft.Row(
                                    [
                                        ft.Container(
                                            content=ft.Text("Час", size=16, weight="bold", color=ft.colors.WHITE),
                                            width=120,
                                            padding=10,
                                        ),
                                        ft.Container(
                                            content=ft.Text("Предмет", size=16, weight="bold", color=ft.colors.WHITE),
                                            expand=True,
                                            padding=10,
                                        ),
                                        ft.Container(
                                            content=ft.Text("Викладач", size=16, weight="bold", color=ft.colors.WHITE),
                                            width=200,
                                            padding=10,
                                        ),
                                        ft.Container(
                                            content=ft.Text("Під'єднатись", size=16, weight="bold", color=ft.colors.WHITE),
                                            width=150,
                                            padding=10,
                                        ),
                                    ],
                                    spacing=0,
                                ),
                                bgcolor=colors["BG_CARD"],
                                border=ft.Border(
                                    top=ft.BorderSide(1, colors["BORDER_COLOR"]),
                                    bottom=ft.BorderSide(1, colors["BORDER_COLOR"]),
                                    left=ft.BorderSide(1, colors["BORDER_COLOR"]),
                                    right=ft.BorderSide(1, colors["BORDER_COLOR"]),
                                ),
                            )
                        )

                        for time_idx, (time_str, subject_infos) in enumerate(time_entries):

                            subjects_list = []
                            teachers_list = []
                            meet_links_list = []

                            for info in subject_infos:
                                subj = info.get('subject', '')
                                teach = info.get('teacher', '')
                                subjects_list.append(subj)
                                teachers_list.append(teach if teach else "-")

                                meet_link = get_meet_link(subj)
                                meet_links_list.append(meet_link)

                            if len(subjects_list) > 1:
                                subject_str = "\n".join(subjects_list)
                                teacher_str = "\n".join(teachers_list)

                                has_any_link = any(meet_links_list)
                            else:
                                subject_str = subjects_list[0] if subjects_list else ""
                                teacher_str = teachers_list[0] if teachers_list else "-"
                                has_any_link = meet_links_list[0] is not None if meet_links_list else False

                            if has_any_link:
                                if len(subjects_list) > 1:

                                    connect_buttons = []
                                    for subj, link in zip(subjects_list, meet_links_list):
                                        if link:

                                            def make_open_handler(url):
                                                return lambda e: webbrowser.open(url)
                                            connect_buttons.append(
                                                ft.TextButton(
                                                    "Під'єднатись",
                                                    on_click=make_open_handler(link),
                                                    style=ft.ButtonStyle(color=ft.colors.BLUE),
                                                )
                                            )
                                    connect_widget = ft.Column(connect_buttons, spacing=2, horizontal_alignment=ft.CrossAxisAlignment.CENTER) if connect_buttons else ft.Container()
                                else:

                                    meet_link = meet_links_list[0]
                                    def make_open_handler(url):
                                        return lambda e: webbrowser.open(url)
                                    connect_widget = ft.TextButton(
                                        "Під'єднатись",
                                        on_click=make_open_handler(meet_link),
                                        style=ft.ButtonStyle(color=ft.colors.BLUE),
                                    )
                            else:
                                connect_widget = ft.Container()

                            row_bg = colors["BG_DARK"] if time_idx % 2 == 0 else colors["BG_CARD"]
                            content_list.append(
                                ft.Container(
                                    content=ft.Row(
                                        [
                                            ft.Container(
                                                content=ft.Text(time_str, size=14, color=ft.colors.WHITE, weight="bold"),
                                                width=120,
                                                padding=10,
                                                alignment=ft.Alignment(0, 0),
                                            ),
                                            ft.Container(
                                                content=ft.Text(subject_str, size=14, color=ft.colors.WHITE),
                                                expand=True,
                                                padding=10,
                                                alignment=ft.Alignment(-1, 0),
                                            ),
                                            ft.Container(
                                                content=ft.Text(teacher_str, size=14, color=ft.colors.WHITE),
                                                width=200,
                                                padding=10,
                                                alignment=ft.Alignment(-1, 0),
                                            ),
                                            ft.Container(
                                                content=connect_widget,
                                                width=150,
                                                padding=10,
                                                alignment=ft.Alignment(0, 0),
                                            ),
                                        ],
                                        spacing=0,
                                    ),
                                    bgcolor=row_bg,
                                    border=ft.Border(
                                        bottom=ft.BorderSide(1, colors["BORDER_COLOR"]),
                                        left=ft.BorderSide(1, colors["BORDER_COLOR"]),
                                        right=ft.BorderSide(1, colors["BORDER_COLOR"]),
                                    ),
                                )
                            )

                        if day_idx < len(days_schedule) - 1:
                            content_list.append(ft.Container(height=20))

            return ft.Container(
                padding=20,
                content=ft.Column(
                    content_list,
                    spacing=5,
                    alignment=ft.MainAxisAlignment.START,
                    horizontal_alignment=ft.CrossAxisAlignment.START,
                    scroll=ft.ScrollMode.AUTO,
                    expand=True,
                ),
            )

        if state["viewing_images"] is not None:
            return ft.Container(
                content=ft.Column(
                    [
                        ft.Row(
                            [
                                ft.IconButton(
                                    icon=ft.icons.ARROW_BACK,
                                    icon_color=ft.colors.WHITE,
                                    on_click=close_viewer,
                                    tooltip="Назад",
                                ),
                                ft.Text(
                                    f"Перегляд: {state.get('viewing_file_name', 'Розклад')}",
                                    size=20,
                                    weight="bold"
                                ),
                            ],
                            alignment=ft.MainAxisAlignment.START,
                        ),
                        ft.Divider(thickness=1, color=colors["BORDER_COLOR"]),
                        ft.Container(
                            content=ft.Column(
                                [ft.Image(src=image, fit=ft.ImageFit.CONTAIN) for image in state["viewing_images"]],
                                scroll=ft.ScrollMode.AUTO,
                                expand=True,
                            ),
                            expand=True,
                        ),
                    ],
                    expand=True,
                ),
                padding=24,
            )

        if "schedule_loading" not in state:
            state["schedule_loading"] = False

        if not state["schedule_files"] and not state["schedule_loading"]:
            state["schedule_loading"] = True

            async def load_schedule():
                state["schedule_error"] = ""
                try:
                    p = await asyncio.to_thread(get_parser)
                    if p is None:
                        state["schedule_files"] = []
                        state["schedule_error"] = "Не вдалося підключитися до lib.istu.edu.ua. Перевірте інтернет."
                    else:
                        state["schedule_url"] = p.base_url + DEFAULT_SCHEDULE_PATH
                        state["schedule_files"] = await asyncio.to_thread(p.parse_files, state["schedule_url"])
                        if not state["schedule_files"]:
                            await asyncio.sleep(0.3)
                            state["schedule_files"] = await asyncio.to_thread(p.parse_files, state["schedule_url"])
                        if not state["schedule_files"]:
                            reset_parser()
                            p2 = await asyncio.to_thread(get_parser)
                            if p2:
                                state["schedule_url"] = p2.base_url + DEFAULT_SCHEDULE_PATH
                                state["schedule_files"] = await asyncio.to_thread(
                                    p2.parse_files, state["schedule_url"]
                                )
                except (requests.exceptions.ConnectionError, OSError) as e:
                    err = str(e).strip() or type(e).__name__
                    state["schedule_error"] = f"Помилка з'єднання: {err}. Перевірте інтернет або VPN."
                    state["schedule_files"] = []
                except requests.exceptions.Timeout:
                    state["schedule_error"] = "Час очікування вийшов. Сайт не відповідає. Спробуйте пізніше."
                    state["schedule_files"] = []
                except Exception as e:
                    state["schedule_error"] = f"Помилка: {type(e).__name__} — {str(e)[:200]}"
                    state["schedule_files"] = []
                finally:
                    state["schedule_loading"] = False
                    update_content()

            page.run_task(load_schedule)

        def retry_schedule(e):
            reset_parser()
            state["schedule_files"] = []
            state["schedule_loading"] = False
            state["schedule_error"] = ""
            try:
                clear_schedule_cache()
            except Exception:
                pass
            update_content()

        title_row = []
        if state["schedule_stack"]:
            title_row.append(
                ft.IconButton(
                    icon=ft.icons.ARROW_BACK,
                    icon_color=ft.colors.WHITE,
                    on_click=go_back_in_schedule,
                    tooltip="Назад",
                )
            )
        else:
            title_row.append(ft.Text("Розклад МНТУ", size=20, weight="bold"))

        def get_file_tiles():
            tiles: list[ft.Control] = []
            if state.get("schedule_loading", False):
                tiles.append(
                    ft.Row(
                        [
                            ft.ProgressRing(width=40, height=40, stroke_width=3),
                            ft.Text("Завантаження розкладу...", size=16, color=colors["TEXT_COLOR"]),
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                        spacing=20,
                    )
                )
                return tiles

            if not state["schedule_files"]:
                err_msg = state.get("schedule_error", "").strip()
                if err_msg:
                    tiles.append(
                        ft.Text(err_msg, size=14, color=ft.colors.AMBER_200)
                    )
                else:
                    tiles.append(
                        ft.Text(
                            "Сайт lib.istu.edu.ua недоступний або розклад ще не завантажено. Перевірте інтернет (пароль: powerpower).",
                            size=14,
                            color=ft.colors.AMBER_200,
                        )
                    )
                tiles.append(
                    ft.Button(
                        "Повторити",
                        icon=ft.icons.REFRESH,
                        on_click=retry_schedule,
                        style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=20, vertical=12), shape=ft.RoundedRectangleBorder(radius=10)),
                    )
                )
                return tiles
            for f in state["schedule_files"]:

                if f.name == "Розклад занять":
                    continue
                icon = ft.icons.FOLDER if f.is_folder else ft.icons.INSERT_DRIVE_FILE
                tiles.append(
                    ft.Container(
                        content=ft.ListTile(
                            leading=ft.Icon(icon, color=colors["ACCENT"], size=22),
                            title=ft.Text(f.name, size=16, color=colors["TEXT_COLOR"]),
                            on_click=lambda e, file=f: on_schedule_item_click(file),
                        ),
                        border_radius=12,
                        bgcolor=colors["BG_CARD"],
                        padding=_padding_symmetric(horizontal=8, vertical=2),
                    )
                )
            return tiles

        return ft.Container(
            content=ft.Column(
                [
                    ft.Row(title_row, alignment=ft.MainAxisAlignment.START),
                    ft.Divider(thickness=1, color=colors["BORDER_COLOR"]),
                    ft.Column(get_file_tiles(), spacing=8, expand=True),
                ],
                expand=True,
            ),
            padding=24,
        )

    def on_schedule_item_click(file: File):
        if file.is_folder:
            p = get_parser()
            if p is None:
                show_snack(page, "Сайт lib.istu.edu.ua недоступний. Перевірте інтернет.")
                return
            state["schedule_stack"].append(state["schedule_url"])
            state["schedule_url"] = file.url
            try:
                state["schedule_files"] = p.parse_files(state["schedule_url"])
                if not state["schedule_files"]:
                    reset_parser()
                    p2 = get_parser()
                    if p2:
                        state["schedule_files"] = p2.parse_files(state["schedule_url"])
            except Exception as ex:
                show_snack(page, str(ex))
                state["schedule_files"] = []
            update_content()
        else:
            async def load_content():

                state["schedule_loading"] = True
                update_content()

                try:

                    is_docx = file.url.lower().endswith(".docx") or file.name.lower().endswith(".docx")

                    if is_docx:
                        docx_bytes = None
                        p = get_parser()
                        if not p:
                            show_snack(
                                page,
                                "Сайт lib.istu.edu.ua недоступний. Натисніть «Повторити» у списку розкладу.",
                            )
                        else:
                            try:
                                response = p.session.get(file.url, timeout=REQUEST_TIMEOUT)
                                response.raise_for_status()
                                docx_bytes = response.content
                            except Exception as e:
                                print(f"[ERROR] Не вдалося завантажити docx: {e}")
                                show_snack(page, "Не вдалося завантажити файл розкладу.")
                                return

                        if docx_bytes:

                            from docx import Document
                            import io

                            doc = Document(io.BytesIO(docx_bytes))
                            all_tables = doc.tables

                            if len(all_tables) >= 1:

                                today = date.today()


                                valid_tables = [table for table in all_tables if len(table.rows) >= 2]


                                week_bucket: dict[date, dict[str, list]] = defaultdict(lambda: defaultdict(list))
                                all_groups: list[str] = []


                                for table in valid_tables:
                                    if len(table.rows) < 2:
                                        continue

                                    header_cells = table.rows[0].cells

                                    date_col_idx = 0
                                    time_col_idx = 1

                                    group_columns_map: dict[int, list[str]] = {}
                                    skip_headers = {"час", "дата", "день", "пари", "пара", "пары", "дисципліна", "викладач", "ауд"}
                                    for col_idx, cell in enumerate(header_cells):
                                        hdr = cell.text.strip()

                                        if col_idx == date_col_idx or col_idx == time_col_idx:
                                            continue

                                        if not hdr:
                                            continue
                                        hdr_lower = hdr.lower()
                                        if hdr_lower in skip_headers:
                                            continue
                                        normalized_hdr_groups = normalize_group_names([hdr])
                                        if not normalized_hdr_groups:
                                            continue
                                        group_columns_map[col_idx] = normalized_hdr_groups
                                        for ng in normalized_hdr_groups:
                                            if ng not in all_groups:
                                                all_groups.append(ng)


                                    last_date_str_raw = ""
                                    for row in table.rows[1:]:
                                        if len(row.cells) < 2:
                                            continue

                                        raw = ""
                                        if date_col_idx < len(row.cells):
                                            raw = row.cells[date_col_idx].text.strip()

                                        if raw and re.search(r"\d", raw) and re.search(r"\d{1,2}[./]\d{1,2}", raw):
                                            date_str_raw = raw
                                            last_date_str_raw = raw
                                        elif last_date_str_raw:
                                            date_str_raw = last_date_str_raw
                                        else:
                                            continue

                                        if not date_str_raw or not re.search(r"\d", date_str_raw):
                                            continue

                                        if not re.search(r"\d{1,2}[./]\d{1,2}", date_str_raw):
                                            continue

                                        parsed_date = extract_date_from_text(date_str_raw, today)

                                        if parsed_date:
                                            date_str = f"{parsed_date.day:02d}.{parsed_date.month:02d}.{parsed_date.year % 100:02d}"
                                        else:
                                            date_match = re.search(r'(\d{1,2})\.(\d{1,2})\.(\d{2,4})', date_str_raw)
                                            if date_match:
                                                date_str = f"{date_match.group(1)}.{date_match.group(2)}.{date_match.group(3)}"
                                            else:
                                                continue

                                        date_str = re.sub(r'\s+', ' ', date_str).strip()

                                        time_str = ""
                                        if time_col_idx < len(row.cells):
                                            time_str = row.cells[time_col_idx].text.strip()

                                        if not time_str:
                                            continue

                                        time_str = re.sub(r'\s+', '', time_str)

                                        if len(time_str) == 4 and time_str.isdigit():
                                            time_str = f"{time_str[:2]}:{time_str[2:]}"

                                        elif len(time_str) == 4 and ":" in time_str:
                                            parts = time_str.split(":")
                                            if len(parts) == 2 and len(parts[0]) == 1:
                                                time_str = f"0{parts[0]}:{parts[1]}"

                                        for col_idx, group_names in group_columns_map.items():
                                            if col_idx >= len(row.cells):
                                                continue

                                            text = row.cells[col_idx].text.strip()

                                            text = re.sub(r'[ \t]+', ' ', text)
                                            text = re.sub(r'\n\s*\n', '\n', text)
                                            text = text.strip()

                                            if not text or text.lower() in ["-", "—", "немає", "відсутній", "пусто", ""]:
                                                continue

                                            tl = re.sub(r'[\s,;.\-_]+', ' ', text.lower()).strip()
                                            if "дисципліна" in tl and "викладач" in tl and ("ауд" in tl or "аудиторія" in tl):
                                                continue
                                            header_only = {"дисципліна", "викладач", "ауд", "аудиторія"}
                                            words = set(w.strip() for w in re.split(r'[\s,;]+', tl) if len(w.strip()) > 1)
                                            if words and words <= header_only:
                                                continue

                                            subject, teacher = extract_teacher_from_text(text)

                                            for group_name in group_names:
                                                entry = {
                                                    "date": date_str,
                                                    "time": time_str,
                                                    "subject": subject,
                                                    "teacher": teacher,
                                                }
                                                pd_for_week = parsed_date or parse_schedule_date_str(date_str)
                                                if pd_for_week:
                                                    monday = pd_for_week - timedelta(days=pd_for_week.weekday())
                                                    week_bucket[monday][group_name].append(entry)

                                sorted_mondays = sorted(week_bucket.keys())
                                if not sorted_mondays:
                                    msg = (
                                        "У таблиці не знайдено пар з датою та часом. "
                                        "Можливо, змінився формат DOCX — повідомте авторів або спробуйте інший файл."
                                    )
                                    if not all_groups:
                                        msg = "У таблиці не знайдено колонок груп (рядок заголовка). Перевірте файл розкладу."
                                    show_snack(page, msg)
                                    state["schedule_loading"] = False
                                    update_content()
                                    return

                                mon_cur = today - timedelta(days=today.weekday())
                                mon_next = mon_cur + timedelta(days=7)
                                week_options = ["Цей тиждень", "Наступний тиждень"]
                                combined: dict[str, dict[str, list]] = {}
                                dates_by_week: dict[str, list[str]] = {"Цей тиждень": [], "Наступний тиждень": []}

                                def _sort_entries(entries: list) -> list:
                                    return sorted(
                                        entries,
                                        key=lambda e: (
                                            parse_schedule_date_str(e.get("date") or "") or date.min,
                                            e.get("time") or "",
                                        ),
                                    )

                                for g in all_groups:
                                    cur = _sort_entries(list(week_bucket.get(mon_cur, {}).get(g, [])))
                                    nex = _sort_entries(list(week_bucket.get(mon_next, {}).get(g, [])))
                                    if not nex:
                                        nex = [
                                            {
                                                "date": "",
                                                "time": "",
                                                "subject": "Інформації немає",
                                                "teacher": "",
                                            }
                                        ]
                                    combined[g] = {"Цей тиждень": cur, "Наступний тиждень": nex}

                                for label, _ in (
                                    ("Цей тиждень", mon_cur),
                                    ("Наступний тиждень", mon_next),
                                ):
                                    acc: set[str] = set()
                                    for g in all_groups:
                                        for e in combined[g][label]:
                                            if _is_synthetic_no_classes_entry(e):
                                                continue
                                            ds = (e.get("date") or "").strip()
                                            if ds:
                                                acc.add(ds)
                                    dates_by_week[label] = sorted(
                                        acc, key=lambda s: parse_schedule_date_str(s) or date.min
                                    )

                                state["parsed_schedule_per_group"] = combined
                                state["group_options"] = normalize_group_names(all_groups)
                                state["week_options"] = week_options
                                state["selected_group"] = ""
                                state["selected_week"] = "Цей тиждень"
                                state["viewing_images"] = None
                                state["week_dates"] = dates_by_week

                                try:
                                    save_schedule_to_db(combined, normalize_group_names(all_groups), week_options, state["week_dates"])
                                except Exception as e:
                                    print(f"[WARN] Не вдалося зберегти розклад: {e}")

                                update_content()
                                return

                    images, docx_bytes = await get_full_file_content(file.url)

                    if images:

                        state["viewing_images"] = images
                        state["viewing_file_name"] = file.name
                        update_content()
                        return

                    if docx_bytes is not None and not is_docx:

                        images, docx_bytes = await get_last_2_pages_content(file.url)
                        if images:
                            state["viewing_images"] = images
                            state["viewing_file_name"] = file.name
                            update_content()
                            return

                    if not images and (docx_bytes is None or is_docx):
                        show_snack(page, "Не вдалося відкрити розклад.")
                    elif images:
                        state["viewing_images"] = images
                        state["viewing_file_name"] = file.name
                        update_content()
                    else:
                        show_snack(page, "Не вдалося відкрити розклад.")
                except Exception as e:
                    print(f"[ERROR] Помилка обробки файлу: {e}")
                    show_snack(page, f"Помилка відкриття файлу: {str(e)[:100]}")
                finally:

                    state["schedule_loading"] = False
                    update_content()

            page.run_task(load_content)

    def go_back_in_schedule(e):
        if state["schedule_stack"]:
            p = get_parser()
            if p is None:
                show_snack(page, "Сайт lib.istu.edu.ua недоступний.")
                return
            prev_url = state["schedule_stack"].pop()
            state["schedule_url"] = prev_url
            try:
                state["schedule_files"] = p.parse_files(state["schedule_url"])
                if not state["schedule_files"]:
                    reset_parser()
                    p2 = get_parser()
                    if p2:
                        state["schedule_files"] = p2.parse_files(state["schedule_url"])
            except Exception as ex:
                show_snack(page, str(ex))
                state["schedule_files"] = []
            update_content()

    def close_group_view(e):
        state["parsed_schedule_per_group"] = {}
        state["group_options"] = []
        state["week_options"] = []
        state["selected_group"] = ""
        state["selected_week"] = ""
        state["week_dates"] = {}
        try:
            clear_schedule_cache()
        except Exception:
            pass
        update_content()

    def close_viewer(e):
        state["viewing_images"] = None
        state["viewing_file_name"] = None
        state["parsed_schedule_per_group"] = {}
        state["group_options"] = []
        state["week_options"] = []
        state["selected_group"] = ""
        state["selected_week"] = ""
        state["week_dates"] = {}
        try:
            clear_schedule_cache()
        except Exception:
            pass
        update_content()

    def select_group(e):
        state["selected_group"] = _dropdown_value_from_event(e) or ""
        update_content()

    def select_week(e):
        state["selected_week"] = _dropdown_value_from_event(e) or ""
        update_content()

    def calendar_page():
        vy = state["visible_year"]
        vm = state["visible_month"]
        sd = state["selected_date"]

        date_header = ft.Text(
            f"{sd.day} {UKR_MONTH_NAMES[sd.month - 1]} {sd.year}",
            size=20,
            weight="bold",
            text_align=ft.TextAlign.CENTER,
        )
        month_label = ft.Text(f"{UKR_MONTH_NAMES[vm - 1]} {vy}", size=18, weight="bold", text_align=ft.TextAlign.CENTER)
        footer_label = ft.Text(f"Обрано: {sd.day}.{sd.month}.{sd.year}", size=14)

        def update_month_label():
            vy2 = state["visible_year"]
            vm2 = state["visible_month"]
            month_label.value = f"{UKR_MONTH_NAMES[vm2 - 1]} {vy2}"
            month_label.update()

        def update_date_header():
            ss = state["selected_date"]
            date_header.value = f"{ss.day} {UKR_MONTH_NAMES[ss.month - 1]} {ss.year}"
            date_header.update()

        def update_footer_label():
            ss = state["selected_date"]
            footer_label.value = f"Обрано: {ss.day}.{ss.month}.{ss.year}"
            footer_label.update()

        def go_previous_month(e):
            m = state["visible_month"] - 1
            y = state["visible_year"]
            if m == 0:
                m = 12
                y -= 1
            state["visible_month"] = m
            state["visible_year"] = y
            refresh_grid()
            update_month_label()

        def go_next_month(e):
            m = state["visible_month"] + 1
            y = state["visible_year"]
            if m == 13:
                m = 1
                y += 1
            state["visible_month"] = m
            state["visible_year"] = y
            refresh_grid()
            update_month_label()

        def return_to_today(e):
            state["selected_date"] = today
            state["visible_month"] = today.month
            state["visible_year"] = today.year
            refresh_grid()
            update_month_label()
            update_date_header()
            update_footer_label()

        def pick_date_manually(e):
            init_day = state["selected_date"].day
            init_month = state["selected_date"].month
            init_year = state["selected_date"].year

            def days_in_month(year, month):
                return calendar.monthrange(year, month)[1]

            day_dd = ft.Dropdown(
                width=100,
                value=str(init_day),
                options=[ft.dropdown.Option(str(d)) for d in range(1, days_in_month(init_year, init_month) + 1)],
                **_dropdown_menu_height_kw(160),
            )
            month_dd = ft.Dropdown(
                width=150,
                value=str(init_month),
                options=[ft.dropdown.Option(str(i + 1), UKR_MONTH_NAMES[i]) for i in range(12)],
                **_dropdown_menu_height_kw(160),
            )
            year_range = list(range(today.year - 5, 2101))
            year_dd = ft.Dropdown(
                width=130,
                value=str(init_year),
                options=[ft.dropdown.Option(str(y)) for y in year_range],
                **_dropdown_menu_height_kw(160),
            )

            def update_days(e):
                y = int(year_dd.value)
                m = int(month_dd.value)
                max_day = days_in_month(y, m)
                day_dd.options = [ft.dropdown.Option(str(d)) for d in range(1, max_day + 1)]
                if int(day_dd.value) > max_day:
                    day_dd.value = str(max_day)
                day_dd.update()

            _bind_dropdown_change(month_dd, update_days)
            _bind_dropdown_change(year_dd, update_days)

            def on_ok(ev):
                chosen_day = int(day_dd.value)
                chosen_month = int(month_dd.value)
                chosen_year = int(year_dd.value)
                state["selected_date"] = date(chosen_year, chosen_month, chosen_day)
                state["visible_month"] = chosen_month
                state["visible_year"] = chosen_year
                refresh_grid()
                update_month_label()
                update_date_header()
                update_footer_label()
                dialog.open = False
                page.update()

            def on_cancel(ev):
                dialog.open = False
                page.update()

            nonlocal dialog
            dialog = ft.AlertDialog(
                title=ft.Container(ft.Text("Оберіть дату", size=16, weight="bold"), alignment=ft.Alignment(0, 0)),
                content=ft.Container(
                    ft.Column(
                        [
                            ft.Container(
                                ft.Row(
                                    [
                                        ft.Text("День:"),
                                        day_dd,
                                        ft.Container(width=20),
                                        ft.Text("Місяць:"),
                                        month_dd,
                                        ft.Container(width=20),
                                        ft.Text("Рік:"),
                                        year_dd,
                                    ],
                                    alignment=ft.MainAxisAlignment.START,
                                    spacing=10,
                                ),
                                margin=ft.margin.only(left=-10),
                            )
                        ]
                    ),
                    padding=_padding_all(10),
                    width=600,
                    height=220,
                ),
                actions=[
                    ft.Button(
                        "OK",
                        on_click=on_ok,
                        bgcolor=colors["ACCENT"],
                        style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
                    ),
                    ft.Button(
                        "Скасувати",
                        on_click=on_cancel,
                        bgcolor=ft.colors.RED_600,
                        style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
                    ),
                ],
                actions_alignment=ft.MainAxisAlignment.CENTER,
            )
            page.overlay.append(dialog)
            dialog.open = True
            page.update()

        CALENDAR_WIDTH = 7 * 80 + 6 * 5
        grid = ft.GridView(
            expand=True,
            max_extent=80,
            child_aspect_ratio=1,
            spacing=5,
        )

        last_click_time = {}
        last_click_day = {}

        def show_day_details(target_date: date, entries: list[dict], group_name: str):
            if not entries:
                show_snack(page, f"На {target_date.strftime('%d.%m.%Y')} пар немає.")
                return

            date_str = target_date.strftime("%d.%m.%Y")
            day_name = ["Понеділок", "Вівторок", "Середа", "Четвер", "П'ятниця", "Субота", "Неділя"][target_date.weekday()]

            content_list = []
            content_list.append(ft.Text(f"{day_name}, {date_str}", size=18, weight="bold", color=ft.colors.WHITE))
            content_list.append(ft.Divider(thickness=1, color=colors["BORDER_COLOR"]))

            time_groups = {}
            for entry in entries:
                time_key = entry.get('time', '')
                if time_key not in time_groups:
                    time_groups[time_key] = []
                time_groups[time_key].append(entry)

            sorted_times = sorted(time_groups.keys())

            for time_str in sorted_times:
                subject_infos = time_groups[time_str]
                for info in subject_infos:
                    subject = info.get('subject', '')
                    teacher = info.get('teacher', '')
                    meet_link = get_meet_link(subject)

                    content_list.append(ft.Container(height=8))
                    content_list.append(ft.Text(f"Час: {time_str}", size=14, weight="bold", color=ft.colors.WHITE))
                    content_list.append(ft.Text(f"Предмет: {subject}", size=14, color=ft.colors.WHITE))
                    content_list.append(ft.Text(f"Викладач: {teacher if teacher else '-'}", size=14, color=ft.colors.WHITE))
                    if meet_link:
                        def make_open_handler(url):
                            return lambda e: webbrowser.open(url)
                        content_list.append(
                            ft.TextButton(
                                "Під'єднатись до Google Meet",
                                on_click=make_open_handler(meet_link),
                                style=ft.ButtonStyle(color=ft.colors.BLUE),
                            )
                        )
                    content_list.append(ft.Divider(thickness=1, color=colors["BORDER_COLOR"]))

            dialog = ft.AlertDialog(
                title=ft.Text(f"Розклад на {date_str}", size=18, weight="bold"),
                content=ft.Container(
                    content=ft.Column(
                        content_list,
                        spacing=5,
                        scroll=ft.ScrollMode.AUTO,
                    ),
                    width=500,
                    height=400,
                    padding=20,
                ),
                actions=[
                    ft.Button(
                        "Закрити",
                        on_click=lambda e: close_day_details_dialog(dialog),
                        bgcolor=colors["ACCENT"],
                        style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
                    )
                ],
                actions_alignment=ft.MainAxisAlignment.CENTER,
            )
            page.dialog = dialog
            dialog.open = True
            page.update()

        def close_day_details_dialog(dialog: ft.AlertDialog):
            dialog.open = False
            page.update()

        def get_subjects_for_date(target_date: date, group_name: str) -> tuple[list[str], list[dict]]:
            if not group_name or not state.get("parsed_schedule_per_group"):
                return [], []

            if group_name not in state["parsed_schedule_per_group"]:
                return [], []

            _, group_schedule = get_group_schedule(group_name)
            date_str_2digit = target_date.strftime("%d.%m.%y")
            date_str_4digit = target_date.strftime("%d.%m.%Y")
            date_str_no_leading = f"{target_date.day}.{target_date.month}.{target_date.year % 100}"

            subjects = []
            full_entries = []
            for week_key, entries in group_schedule.items():
                for entry in entries:
                    entry_date = entry.get("date", "").strip()
                    if entry_date in [date_str_2digit, date_str_4digit, date_str_no_leading]:
                        subject = entry.get("subject", "").strip()
                        if subject:
                            if subject not in subjects:
                                subjects.append(subject)
                            full_entries.append(entry)

            return subjects, full_entries

        def refresh_grid():
            grid.controls.clear()
            vy2 = state["visible_year"]
            vm2 = state["visible_month"]
            sd2 = state["selected_date"]

            selected_group = state.get("notify_group") or state.get("selected_group", "")

            month_matrix = calendar.monthcalendar(vy2, vm2)

            for week in month_matrix:
                for day_num in week:
                    if day_num == 0:
                        cell = ft.Container(
                            ft.Text(""),
                            bgcolor=colors["BG_DARK"],
                            border=_border_all(1, colors["BORDER_COLOR"]),
                            alignment=ft.Alignment(0, 0),
                        )
                    else:
                        is_selected = (vy2 == sd2.year and vm2 == sd2.month and day_num == sd2.day)
                        is_today = (vy2 == today.year and vm2 == today.month and day_num == today.day)

                        if is_selected:
                            bg = colors["ACCENT"]
                            border_color = colors["ACCENT"]
                        elif is_today:
                            bg = colors["BG_CARD"]
                            border_color = colors["ACCENT"]
                        else:
                            bg = colors["BG_CARD"]
                            border_color = colors["BORDER_COLOR"]

                        current_date = date(vy2, vm2, day_num)
                        subjects, full_entries = get_subjects_for_date(current_date, selected_group) if selected_group else ([], [])

                        cell_content = []
                        cell_content.append(ft.Text(str(day_num), color=ft.colors.WHITE, size=14, weight="bold"))

                        if subjects:
                            cell_content.append(ft.Container(height=2))
                            for i, subj in enumerate(subjects[:2]):
                                subject_short = subj[:15] + "..." if len(subj) > 15 else subj
                                cell_content.append(
                                    ft.Text(
                                        subject_short,
                                        color=ft.colors.WHITE,
                                        size=8,
                                        overflow=ft.TextOverflow.ELLIPSIS,
                                    )
                                )
                            if len(subjects) > 2:
                                cell_content.append(ft.Text(f"+{len(subjects) - 2}", size=8, color=ft.colors.WHITE))

                        def make_click(day=day_num, entries=full_entries):
                            def on_click(ev):
                                current_time = time.time()
                                day_key = f"{vy2}-{vm2}-{day}"
                                target_date = date(vy2, vm2, day)

                                if day_key in last_click_time and current_time - last_click_time[day_key] < 0.5:
                                    if last_click_day.get(day_key) == day:
                                        switch_to_schedule_with_date(target_date)
                                        last_click_time.pop(day_key, None)
                                        last_click_day.pop(day_key, None)
                                        return

                                last_click_time[day_key] = current_time
                                last_click_day[day_key] = day

                                state["selected_date"] = target_date
                                update_date_header()
                                update_footer_label()
                                refresh_grid()
                            return on_click

                        border_width = 2 if is_today else 1
                        cell = ft.Container(
                            content=ft.Column(
                                cell_content,
                                spacing=1,
                                tight=True,
                                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                            ),
                            bgcolor=bg,
                            border=_border_all(border_width, border_color),
                            border_radius=10,
                            alignment=ft.Alignment(0, -1),
                            padding=_padding_all(4),
                            on_click=make_click(),
                        )
                    grid.controls.append(cell)

            page.update()

        refresh_grid()
        dialog = None

        selected_group = state.get("notify_group") or state.get("selected_group", "")
        group_info_text = ft.Text(
            f"Група: {selected_group}" if selected_group else "Група не вибрана (оберіть в налаштуваннях)",
            size=14,
            color=colors["TEXT_MUTED"] if not selected_group else colors["TEXT_COLOR"],
            text_align=ft.TextAlign.CENTER,
        )

        def update_group_info():
            selected_group_new = state.get("notify_group") or state.get("selected_group", "")
            group_info_text.value = f"Група: {selected_group_new}" if selected_group_new else "Група не вибрана (оберіть в налаштуваннях)"
            group_info_text.color = colors["TEXT_MUTED"] if not selected_group_new else colors["TEXT_COLOR"]
            group_info_text.update()
            refresh_grid()

        return ft.Container(
            content=ft.Column(
                [
                    ft.Container(date_header, alignment=ft.Alignment(0, 0)),
                    ft.Container(group_info_text, alignment=ft.Alignment(0, 0)),
                    ft.Container(height=10),
                    ft.Row(
                        [
                            ft.IconButton(icon=ft.icons.ARROW_BACK, icon_color=ft.colors.WHITE, on_click=go_previous_month),
                            month_label,
                            ft.IconButton(icon=ft.icons.ARROW_FORWARD, icon_color=ft.colors.WHITE, on_click=go_next_month),
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                        spacing=10,
                    ),
                    ft.Container(height=10),
                    ft.Row(
                        [
                            ft.Button(
                                "Вибрати дату вручну",
                                on_click=pick_date_manually,
                                bgcolor=colors["ACCENT"],
                                style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=20, vertical=12), shape=ft.RoundedRectangleBorder(radius=10)),
                            ),
                            ft.Container(width=20),
                            ft.Button(
                                "Повернутися до сьогоднішньої дати",
                                on_click=return_to_today,
                                bgcolor=colors["ACCENT"],
                                style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=20, vertical=12), shape=ft.RoundedRectangleBorder(radius=10)),
                            ),
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                    ),
                    ft.Container(height=10),
                    ft.Container(
                        padding=_padding_only(left=24),
                        content=ft.Column(
                            [
                                ft.Row(
                                    [
                                        ft.Container(ft.Row(expand=True, controls=[ft.Container(expand=True), ft.Text("Пн"), ft.Container(expand=True)], alignment=ft.MainAxisAlignment.CENTER), width=80),
                                        ft.Container(ft.Row(expand=True, controls=[ft.Container(expand=True), ft.Text("Вт"), ft.Container(expand=True)], alignment=ft.MainAxisAlignment.CENTER), width=80),
                                        ft.Container(ft.Row(expand=True, controls=[ft.Container(expand=True), ft.Text("Ср"), ft.Container(expand=True)], alignment=ft.MainAxisAlignment.CENTER), width=80),
                                        ft.Container(ft.Row(expand=True, controls=[ft.Container(expand=True), ft.Text("Чт"), ft.Container(expand=True)], alignment=ft.MainAxisAlignment.CENTER), width=80),
                                        ft.Container(ft.Row(expand=True, controls=[ft.Container(expand=True), ft.Text("Пт"), ft.Container(expand=True)], alignment=ft.MainAxisAlignment.CENTER), width=80),
                                        ft.Container(ft.Row(expand=True, controls=[ft.Container(expand=True), ft.Text("Сб"), ft.Container(expand=True)], alignment=ft.MainAxisAlignment.CENTER), width=80),
                                        ft.Container(ft.Row(expand=True, controls=[ft.Container(expand=True), ft.Text("Нд"), ft.Container(expand=True)], alignment=ft.MainAxisAlignment.CENTER), width=80),
                                    ],
                                    spacing=5,
                                    alignment=ft.MainAxisAlignment.START,
                                ),
                                ft.Container(height=5),
                                ft.Container(grid, width=CALENDAR_WIDTH, expand=True),
                            ],
                            horizontal_alignment=ft.CrossAxisAlignment.START,
                        ),
                    ),
                    ft.Container(height=10),
                    footer_label,
                ],
                alignment=ft.MainAxisAlignment.START,
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                expand=True,
            ),
            padding=24,
        )

    def settings_page():
        def apply_theme_refresh():
            page.bgcolor = colors["BG_DARK"]
            content_area.bgcolor = colors["BG_DARK"]
            content_area.content.bgcolor = colors["BG_LIGHT"]
            sidebar.bgcolor = colors["BG_LIGHT"]
            sidebar.content.controls[0].gradient = ft.LinearGradient(
                begin=ft.Alignment(-1, -1),
                end=ft.Alignment(1, 1),
                colors=[colors["BG_CARD"], "#1a1f26"],
            )
            for idx, container in enumerate(nav_buttons):
                container.bgcolor = colors["ACCENT"] if idx == selected_index else "transparent"
            update_content()

        def open_color_picker(e):
            r, g, b = hex_to_rgb(colors["ACCENT"])
            h0, s0, l0 = rgb_to_hsl(r, g, b)
            picker_state = {"h": h0, "s": s0, "l": l0}

            preview = ft.Container(width=140, height=44, border_radius=10, bgcolor=colors["ACCENT"], border=_border_all(2, colors["BORDER_COLOR"]))
            hex_field = ft.TextField(
                label="Колір (#rrggbb) — можна вписати вручну",
                value=colors["ACCENT"],
                width=260,
                on_change=lambda e: on_hex_changed(),
            )
            s_slider = ft.Slider(min=0, max=100, value=s0, divisions=100, label="Насиченість: {value}%", width=260, on_change=lambda e: on_sl_changed())
            l_slider = ft.Slider(min=0, max=100, value=l0, divisions=100, label="Світлість: {value}%", width=260, on_change=lambda e: on_sl_changed())

            apply_button_ref = ft.Ref[ft.Button]()

            def sync_from_hsl():
                hex_val = hsl_to_hex(picker_state["h"], picker_state["s"], picker_state["l"])
                preview.bgcolor = hex_val
                hex_field.value = hex_val

                if apply_button_ref and apply_button_ref.current:
                    apply_button_ref.current.bgcolor = hex_val

            def set_color_from_xy(x: float, y: float, palette_size: float):
                try:

                    hue_norm = max(0, min(1, x / palette_size))
                    sat_norm = 1.0 - max(0, min(1, y / palette_size))

                    hue_deg = hue_norm * 360
                    saturation = sat_norm * 100

                    picker_state["h"] = hue_deg
                    picker_state["s"] = saturation

                    if picker_state["l"] == 0:
                        picker_state["l"] = 50

                    s_slider.value = saturation
                    sync_from_hsl()
                    page.update()
                except Exception:
                    pass

            def on_sl_changed():
                picker_state["s"] = int(s_slider.value or 50)
                picker_state["l"] = int(l_slider.value or 50)
                sync_from_hsl()
                page.update()

            def on_hex_changed():
                raw = (hex_field.value or "").strip().lstrip("#")
                if len(raw) == 6:
                    try:
                        int(raw, 16)
                        rr, gg, bb = hex_to_rgb("#" + raw)
                        picker_state["h"], picker_state["s"], picker_state["l"] = rgb_to_hsl(rr, gg, bb)
                        preview.bgcolor = "#" + raw
                        s_slider.value = picker_state["s"]
                        l_slider.value = picker_state["l"]

                        if apply_button_ref and apply_button_ref.current:
                            apply_button_ref.current.bgcolor = "#" + raw
                    except (ValueError, TypeError):
                        pass
                page.update()

            palette_size = 220
            n_hue_steps = 12

            color_grid = []
            for row in range(n_hue_steps):
                row_colors = []
                for col in range(n_hue_steps):
                    hue = (col / n_hue_steps) * 360
                    sat = 100 - (row / n_hue_steps) * 100
                    color = hsl_to_hex(hue, sat, 50)
                    row_colors.append(color)
                color_grid.append(row_colors)

            palette_cells = []
            cell_size = palette_size / n_hue_steps

            def make_color_cell(row, col, color):
                def on_cell_click(e):
                    hue = (col / n_hue_steps) * 360
                    sat = 100 - (row / n_hue_steps) * 100
                    picker_state["h"] = hue
                    picker_state["s"] = sat
                    s_slider.value = sat
                    sync_from_hsl()
                    page.update()

                return ft.Container(
                    width=cell_size,
                    height=cell_size,
                    bgcolor=color,
                    border=_border_all(1, colors["BORDER_COLOR"]),
                    on_click=on_cell_click,
                )

            for row in range(n_hue_steps):
                for col in range(n_hue_steps):
                    color = color_grid[row][col]
                    palette_cells.append(make_color_cell(row, col, color))

            color_palette = ft.GridView(
                palette_cells,
                runs_count=n_hue_steps,
                max_extent=cell_size,
                spacing=0,
                run_spacing=0,
                width=palette_size,
                height=palette_size,
            )

            color_palette_container = ft.Container(
                content=color_palette,
                width=palette_size,
                height=palette_size,
                border=_border_all(2, colors["BORDER_COLOR"]),
                border_radius=8,
            )

            def on_picker_ok(e):
                raw = (hex_field.value or "").strip().lstrip("#")
                if len(raw) == 6:
                    try:
                        int(raw, 16)
                        hex_full = "#" + raw
                        state["theme"] = "custom"
                        state["custom_accent"] = hex_full
                        colors["ACCENT"] = hex_full
                        colors["ACCENT_TEXT"] = accent_text_color(hex_full)
                        save_custom_accent(hex_full)
                        apply_theme_refresh()
                    except ValueError:
                        pass
                color_picker_dialog.open = False
                page.update()

            def on_picker_cancel(e):
                color_picker_dialog.open = False
                page.update()

            color_picker_dialog = ft.AlertDialog(
                title=ft.Text("Обрати колір"),
                content=ft.Container(
                    content=ft.Column(
                        [
                            color_palette_container,
                            ft.Text("Клікніть на колір для вибору відтінку та насиченості", size=11, color=colors["TEXT_MUTED"], text_align=ft.TextAlign.CENTER),
                            ft.Container(height=8),
                            ft.Text("Насиченість і світлість", size=12, color=colors["TEXT_MUTED"]),
                            s_slider,
                            l_slider,
                            ft.Container(height=8),
                            preview,
                            ft.Container(height=8),
                            hex_field,
                        ],
                        spacing=4,
                        width=280,
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        scroll=ft.ScrollMode.AUTO,
                    ),
                    padding=12,
                    height=520,
                ),
                actions=[
                    ft.TextButton(
                        "Скасувати",
                        on_click=on_picker_cancel,
                        style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=20, vertical=12), shape=ft.RoundedRectangleBorder(radius=10)),
                    ),
                    ft.Button(
                        "Застосувати",
                        ref=apply_button_ref,
                        on_click=on_picker_ok,
                        bgcolor=colors["ACCENT"],
                        style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
                    ),
                ],
            )
            page.overlay.append(color_picker_dialog)
            color_picker_dialog.open = True
            page.update()

        def on_switch_change(e):
            state["notifications_enabled"] = e.control.value
            update_content()

        def on_slider_change(e):
            old_value = state.get("notify_minutes", 120)
            new_value = int(round(e.control.value))
            state["notify_minutes"] = new_value
            notify_label.value = f"Час повідомлення про пару: {state['notify_minutes']} хв"
            notify_label.update()

            if old_value != new_value:
                clear_today_notifications()
                state["notified_entries"] = set()
                print(f"[NOTIFICATION] Час уведомлення змінено з {old_value} на {new_value} хв, очищено показані уведомлення")

        def on_group_change(e):
            state["notify_group"] = _dropdown_value_from_event(e) or ""
            update_content()

        def on_telegram_switch_change(e):
            state["telegram_enabled"] = bool(e.control.value)

        def on_telegram_token_change(e):
            state["telegram_bot_token"] = (e.control.value or "").strip()

        def on_telegram_chat_change(e):
            state["telegram_chat_id"] = (e.control.value or "").strip()

        def save_settings(e):
            save_settings_to_db(
                state["notifications_enabled"],
                state["notify_minutes"],
                state["notify_group"],
                state["theme"],
            )
            persist_bot_settings()
            show_snack(
                page,
                f"Налаштування збережено:\n"
                f"  Сповіщення {'увімкнено' if state['notifications_enabled'] else 'вимкнено'}\n"
                f"  Час: {state['notify_minutes']} хв\n"
                f"  Група: {state['notify_group'] or 'не вибрано'}\n"
                f"  Колір: {colors['ACCENT']}\n"
                f"  Telegram: {'увімкнено' if state['telegram_enabled'] else 'вимкнено'}"
            )

        notify_label = ft.Text(f"Час повідомлення про пару: {state['notify_minutes']} хв")

        controls = [
            ft.Text("Налаштування", size=20, weight="bold", color=colors["TEXT_COLOR"]),
            ft.Container(height=24),
            ft.Container(
                content=ft.Column(
                    [
                        ft.Text("Колір інтерфейсу", size=16, weight="bold", color=colors["TEXT_COLOR"]),
                        ft.Text("Натисніть палітру, щоб обрати акцентний колір", size=13, color=colors["TEXT_MUTED"]),
                        ft.Container(height=12),
                        ft.Container(
                            content=ft.Row(
                                [
                                    ft.Container(
                                        content=ft.Row(
                                            [
                                                ft.Container(width=12, height=28, border_radius=4, bgcolor="#f778ba"),
                                                ft.Container(width=12, height=28, border_radius=4, bgcolor="#58a6ff"),
                                                ft.Container(width=12, height=28, border_radius=4, bgcolor="#2ea043"),
                                                ft.Container(width=12, height=28, border_radius=4, bgcolor="#d29922"),
                                                ft.Container(width=12, height=28, border_radius=4, bgcolor="#a371f7"),
                                                ft.Container(width=12, height=28, border_radius=4, bgcolor="#f85149"),
                                            ],
                                            spacing=3,
                                        ),
                                        padding=6,
                                        border_radius=10,
                                        bgcolor=colors["BG_LIGHT"],
                                        border=_border_all(1, colors["BORDER_COLOR"]),
                                    ),
                                    ft.Text("Свій колір", size=14, weight="w600", color=colors["TEXT_COLOR"]),
                                ],
                                spacing=12,
                                alignment=ft.MainAxisAlignment.CENTER,
                            ),
                            on_click=open_color_picker,
                            border_radius=12,
                            padding=10,
                            border=_border_all(2, colors["BORDER_COLOR"]),
                            ink=True,
                            tooltip="Натисніть, щоб обрати будь-який колір",
                        ),
                        ft.Container(height=8),
                        ft.Text(
                            f"Поточний колір: {colors['ACCENT']}",
                            size=13,
                            color=colors["TEXT_MUTED"],
                        ),
                    ],
                    spacing=4,
                ),
                padding=16,
                border_radius=16,
                bgcolor=colors["BG_CARD"],
                border=_border_all(1, colors["BORDER_COLOR"]),
            ),
            ft.Container(height=24),
            ft.Row(
                [
                    ft.Text("Сповіщення", color=colors["TEXT_COLOR"]),
                    ft.Switch(
                        value=state["notifications_enabled"],
                        on_change=on_switch_change,
                    ),
                ],
                alignment=ft.MainAxisAlignment.START,
                spacing=20,
            ),
            ft.Container(height=20),
            ft.Container(
                content=ft.Column(
                    [
                        ft.Text("Telegram-бот", size=16, weight="bold", color=colors["TEXT_COLOR"]),
                        ft.Text("Керуйте тими ж налаштуваннями через бота: група, час, вимк./увімк. сповіщень.", size=12, color=colors["TEXT_MUTED"]),
                        ft.Container(height=8),
                        ft.Row(
                            [
                                ft.Text("Увімкнути Telegram", color=colors["TEXT_COLOR"]),
                                ft.Switch(value=state.get("telegram_enabled", False), on_change=on_telegram_switch_change),
                            ],
                            alignment=ft.MainAxisAlignment.START,
                            spacing=20,
                        ),
                        ft.Container(height=8),
                        ft.TextField(
                            label="Bot token",
                            value=state.get("telegram_bot_token", ""),
                            password=True,
                            can_reveal_password=True,
                            width=460,
                            on_change=on_telegram_token_change,
                        ),
                        ft.TextField(
                            label="Chat ID (залиште порожнім для автоприв'язки через /start)",
                            value=state.get("telegram_chat_id", ""),
                            width=460,
                            on_change=on_telegram_chat_change,
                        ),
                    ],
                    spacing=6,
                ),
                padding=16,
                border_radius=16,
                bgcolor=colors["BG_CARD"],
                border=_border_all(1, colors["BORDER_COLOR"]),
            ),
            ft.Container(height=20),
        ]

        if state["notifications_enabled"]:
            def test_notifications_now(e):
                notifications = find_upcoming_entries()
                if notifications:
                    for notif_data in notifications:
                        show_class_notification(page, notif_data, colors)
                    show_snack(page, f"Знайдено {len(notifications)} нагадувань про пари")
                else:
                    sg = state.get("notify_group") or state.get("selected_group")
                    if not sg:
                        show_snack(page, "Оберіть групу для сповіщень або відкрийте розклад")
                    _, sg_schedule = get_group_schedule(sg)
                    if not sg_schedule:
                        show_snack(page, "Спочатку відкрийте розклад для цієї групи")
                    else:

                        all_groups_schedule = state.get("parsed_schedule_per_group", {})
                        _, group_schedule = get_group_schedule(sg)
                        notify_time = int(state.get("notify_minutes", 120)) or 1
                        now = datetime.now()
                        debug_info = []
                        total_entries = 0

                        for week_name, entries in group_schedule.items():
                            for entry in entries:
                                if not isinstance(entry, dict):
                                    continue
                                total_entries += 1
                                date_str = entry.get("date", "").strip()
                                time_str = entry.get("time", "").strip()
                                subject = entry.get("subject", "").strip()

                                if not date_str or not time_str:
                                    continue

                                try:

                                    class_date = None
                                    today_date = now.date()

                                    parts = date_str.split(".")
                                    if len(parts) == 3:
                                        day, month, year_str = parts
                                        if len(year_str) == 2:
                                            year_2digit = int(year_str)
                                            current_year_2digit = today_date.year % 100
                                            if year_2digit <= current_year_2digit:
                                                if current_year_2digit - year_2digit > 1:
                                                    full_year = (today_date.year // 100) * 100 + year_2digit
                                                else:
                                                    full_year_prev = (today_date.year // 100 - 1) * 100 + year_2digit
                                                    full_year_curr = (today_date.year // 100) * 100 + year_2digit
                                                    if abs(full_year_prev - today_date.year) < abs(full_year_curr - today_date.year):
                                                        full_year = full_year_prev
                                                    else:
                                                        full_year = full_year_curr
                                            else:
                                                full_year = (today_date.year // 100) * 100 + year_2digit
                                            class_date = date(int(full_year), int(month), int(day))
                                        else:
                                            class_date = datetime.strptime(date_str, "%d.%m.%Y").date()

                                    if not class_date:
                                        debug_info.append(f"Не вдалося розпарсити дату: {date_str}")
                                        continue

                                    if class_date != today_date:
                                        continue

                                    original_time_str = time_str
                                    if "-" in time_str:
                                        time_str = time_str.split("-")[0].strip()

                                    if ":" not in time_str:
                                        debug_info.append(f"Невірний формат часу: {original_time_str}")
                                        continue
                                    time_parts = time_str.split(":")
                                    if len(time_parts) != 2:
                                        debug_info.append(f"Невірний формат часу: {original_time_str}")
                                        continue
                                    hour = int(time_parts[0].strip())
                                    minute = int(time_parts[1].strip())

                                    if hour < 0 or hour > 23 or minute < 0 or minute > 59:
                                        debug_info.append(f"Невірний час: {hour}:{minute}")
                                        continue

                                    class_dt = datetime(class_date.year, class_date.month, class_date.day, hour, minute)
                                    delta = class_dt - now

                                    if delta.total_seconds() > 0:
                                        minutes_until = int(math.ceil(delta.total_seconds() / 60))
                                        status = "✅" if 1 <= minutes_until <= notify_time else "⏰"
                                        debug_info.append(f"{status} {date_str} {time_str}: {subject} (через {minutes_until} хв, потрібно ≤{notify_time})")
                                    else:
                                        debug_info.append(f"⏪ {date_str} {time_str}: {subject} (вже пройшла)")

                                except Exception as ex:
                                    debug_info.append(f"❌ Помилка: {date_str} {time_str} - {str(ex)[:50]}")

                        if debug_info:
                            info_text = f"Сьогодні ({today_date.strftime('%d.%m.%Y')}) знайдено пар: {len(debug_info)}. Найближчі:\n" + "\n".join(debug_info[:5])
                            if len(debug_info) > 5:
                                info_text += f"\n... та ще {len(debug_info) - 5}"
                            show_snack(page, info_text)
                        else:
                            show_snack(page, f"Сьогодні ({today_date.strftime('%d.%m.%Y')}) немає пар у межах {notify_time} хв. Перевірте розклад на сьогодні.")

            controls.extend([
                notify_label,
                ft.Slider(
                    min=1,
                    max=120,
                    divisions=119,
                    label="{value} хв",
                    value=state["notify_minutes"],
                    on_change=on_slider_change,
                    width=400,
                ),
                ft.Container(height=12),
                ft.Button(
                    "Перевірити зараз",
                    on_click=test_notifications_now,
                    bgcolor=colors["ACCENT"],
                    width=200,
                    style=ft.ButtonStyle(
                        padding=_padding_symmetric(horizontal=24, vertical=12),
                        shape=ft.RoundedRectangleBorder(radius=10)
                    ),
                ),
                ft.Container(height=20),
            ])

            if state["group_options"]:
                controls.extend([
                    ft.Text("Група для сповіщень:", size=16),
                    ft.Dropdown(
                        value=state["notify_group"] or None,
                        options=[ft.dropdown.Option(g) for g in state["group_options"]],
                        width=300,
                        **_dropdown_event_kw(lambda e: on_group_change(e)),
                    ),
                    ft.Container(height=20),
                ])
            else:
                controls.append(
                    ft.Text(
                        "Групи ще не спарсені. Відкрийте розклад для вибору групи.",
                        color=colors["BORDER_COLOR"],
                        size=14
                    )
                )

        controls.append(
            ft.Button(
                "Зберегти налаштування",
                on_click=save_settings,
                bgcolor=colors["ACCENT"],
                style=ft.ButtonStyle(padding=_padding_symmetric(horizontal=24, vertical=14), shape=ft.RoundedRectangleBorder(radius=10)),
            )
        )

        return ft.Container(
            content=ft.Column(
                controls,
                spacing=15,
            ),
            padding=24,
        )

    content_area = ft.Container(
        expand=True,
        bgcolor=colors["BG_DARK"],
        border_radius=ft.BorderRadius(24, 0, 0, 0),
        content=ft.Container(
            bgcolor=colors["BG_LIGHT"],
            border_radius=24,
            margin=24,
            padding=0,
            clip_behavior=ft.ClipBehavior.HARD_EDGE,
            shadow=ft.BoxShadow(
                blur_radius=24,
                color="#00000088",
                spread_radius=-2,
                offset=ft.Offset(0, 6),
            ),
            content=ft.Container(
                content=ft.Column([ft.Text("Завантаження...", size=16)], expand=True),
                expand=True,
            ),
        ),
    )

    def update_content():
        if not state.get("ui_ready", False):
            return
        nonlocal selected_index
        if "selected_index" in state:
            selected_index = state["selected_index"]
        else:
            selected_index = 0
        try:
            if selected_index == 0:
                inner = news_page()
            elif selected_index == 1:
                inner = schedule_page()
            elif selected_index == 2:
                inner = calendar_page()
            elif selected_index == 3:
                inner = messages_page()
            elif selected_index == 4:
                inner = settings_page()
            else:
                inner = profile_page()

            content_area.content.content = ft.Container(
                content=ft.Column([inner], scroll=ft.ScrollMode.AUTO, expand=True),
                expand=True,
            )

            for idx, container in enumerate(nav_buttons):
                container.bgcolor = colors["ACCENT"] if idx == selected_index else "transparent"

                if idx == 3:
                    unread_count = len(get_unread_announcements(target_group=current_target_group()))
                    messages_badge_text = None
                    if unread_count > 0:
                        messages_badge_text = f"({unread_count})" if unread_count < 100 else "(99+)"

                    row_content = container.content
                    if isinstance(row_content, ft.Row):
                        badge_exists = len(row_content.controls) > 2 and isinstance(row_content.controls[-1], ft.Container)

                        if messages_badge_text:
                            if badge_exists:
                                badge_container = row_content.controls[-1]
                                if isinstance(badge_container.content, ft.Text):
                                    badge_container.content.value = messages_badge_text
                                    badge_container.content.update()
                            else:
                                badge_container = ft.Container(
                                    content=ft.Text(messages_badge_text, size=11, color=ft.colors.WHITE, weight="bold"),
                                    bgcolor=colors["ACCENT"],
                                    padding=ft.Padding(6, 2, 6, 2),
                                    border_radius=10,
                                )
                                row_content.controls.append(badge_container)
                                row_content.update()
                        else:
                            if badge_exists:
                                row_content.controls.pop()
                                row_content.update()

                container.update()

            page.update()
        except Exception as ex:
            show_snack(page, f"Помилка: {ex}")

    def nav_clicked(e, idx: int):
        nonlocal selected_index
        selected_index = idx
        state["selected_index"] = idx
        update_content()

    def switch_to_schedule_with_date(target_date: date):
        selected_group = state.get("notify_group") or state.get("selected_group", "")
        if not selected_group or not state.get("parsed_schedule_per_group"):
            show_snack(page, "Спочатку оберіть групу в налаштуваннях та відкрийте розклад.")
            return

        date_str_2digit = target_date.strftime("%d.%m.%y")
        date_str_4digit = target_date.strftime("%d.%m.%Y")
        date_str_no_leading = f"{target_date.day}.{target_date.month}.{target_date.year % 100}"

        week_dates = state.get("week_dates", {})
        target_week = None

        for week_name, dates_list in week_dates.items():
            for d in dates_list:
                d_clean = d.strip()
                if d_clean in [date_str_2digit, date_str_4digit, date_str_no_leading]:
                    target_week = week_name
                    break
                parsed_schedule_date = extract_date_from_text(d_clean, target_date)
                if parsed_schedule_date and parsed_schedule_date == target_date:
                    target_week = week_name
                    break
            if target_week:
                break

        if target_week:
            state["selected_week"] = target_week
            state["selected_group"] = selected_group
            state["selected_index"] = 1
            update_content()
            return

        cal_today = date.today()
        mon_cur = cal_today - timedelta(days=cal_today.weekday())
        mon_next = mon_cur + timedelta(days=7)
        if mon_cur <= target_date < mon_cur + timedelta(days=7):
            target_week = "Цей тиждень"
        elif mon_next <= target_date < mon_next + timedelta(days=7):
            target_week = "Наступний тиждень"

        if not target_week:
            want_mon = target_date - timedelta(days=target_date.weekday())
            for wname in state.get("week_options", []):
                for dstr in week_dates.get(wname, []):
                    pd = parse_schedule_date_str(dstr) or extract_date_from_text(dstr, target_date)
                    if not pd:
                        continue
                    w_mon = pd - timedelta(days=pd.weekday())
                    if w_mon == want_mon:
                        target_week = wname
                        break
                if target_week:
                    break

        if target_week:
            state["selected_week"] = target_week
            state["selected_group"] = selected_group
            state["selected_index"] = 1
            update_content()
            return

        show_snack(
            page,
            f"Для {target_date.strftime('%d.%m.%Y')} немає вибраного тижня в завантаженому розкладі. Відкрийте актуальний файл з lib.istu.edu.ua.",
        )

    selected_index = 0
    nav_buttons = []

    def make_nav_item(icon, label, index, badge_text=None):
        is_selected = (selected_index == index)
        row_items = [
            ft.Icon(icon, color=ft.colors.WHITE, size=22),
            ft.Text(label, size=15, color=ft.colors.WHITE, weight="w600", expand=True),
        ]

        if badge_text:
            badge = ft.Container(
                content=ft.Text(badge_text, size=11, color=ft.colors.WHITE, weight="bold"),
                bgcolor=colors["ACCENT"],
                padding=ft.Padding(6, 2, 6, 2),
                border_radius=10,
            )
            row_items.append(badge)

        container = ft.Container(
            on_click=lambda e, i=index: nav_clicked(e, i),
            bgcolor=colors["ACCENT"] if is_selected else "transparent",
            border_radius=12,
            padding=12,
            margin=ft.Margin(12, 3, 12, 3),
            content=ft.Row(
                row_items,
                spacing=12,
            ),
            ink=True,
        )
        nav_buttons.append(container)
        return container

    unread_count = len(get_unread_announcements(target_group=current_target_group()))
    messages_badge = None
    if unread_count > 0:
        messages_badge = f"({unread_count})" if unread_count < 100 else "(99+)"

    sidebar = ft.Container(
        bgcolor=colors["BG_LIGHT"],
        width=260,
        border_radius=ft.BorderRadius(0, 0, 0, 0),
        shadow=ft.BoxShadow(blur_radius=20, color="#00000055", spread_radius=0, offset=ft.Offset(4, 0)),
        border=ft.Border(right=ft.BorderSide(1, colors["BORDER_COLOR"])),
        content=ft.Column(
            [
                ft.Container(
                    ft.Row(
                        [
                            ft.Icon(ft.icons.SCHOOL, color=ft.colors.WHITE, size=30),
                            ft.Text("МНТУ Помічник", size=20, weight="bold", color=ft.colors.WHITE),
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                        spacing=12,
                    ),
                    height=72,
                    padding=ft.Padding(0, 14, 0, 0),
                    gradient=ft.LinearGradient(
                        begin=ft.Alignment(-1, -1),
                        end=ft.Alignment(1, 1),
                        colors=[colors["BG_CARD"], "#1a1f26"],
                    ),
                    border_radius=ft.BorderRadius(0, 0, 0, 0),
                ),
                ft.Divider(height=1, color=colors["BORDER_COLOR"]),
                ft.Container(height=12),
                make_nav_item(ft.icons.FEED_ROUNDED, "Новини", 0),
                make_nav_item(ft.icons.SCHEDULE, "Розклад", 1),
                make_nav_item(ft.icons.CALENDAR_MONTH, "Календар", 2),
                make_nav_item(ft.icons.NOTIFICATIONS, "Повідомлення", 3, badge_text=messages_badge),
                ft.Container(expand=True),
                ft.Divider(color=colors["BORDER_COLOR"]),
                make_nav_item(ft.icons.SETTINGS, "Налаштування", 4),
                ft.Container(height=8),
                make_nav_item(ft.icons.PERSON, "Профіль", 5),
                ft.Container(height=20),
            ],
            expand=True,
        ),
    )

    def profile_page():
        nonlocal selected_index
        user_role = state.get("user_role", "Студент")

        def go_to_tab(idx: int):
            nonlocal selected_index
            selected_index = idx
            update_content()

        def make_action_row(icon_name, label, on_click):
            return ft.Container(
                content=ft.Row(
                    [
                        ft.Icon(icon_name, color=colors["ACCENT"], size=20),
                        ft.Text(label, size=15, color=colors["TEXT_COLOR"], expand=True),
                        ft.Icon(ft.icons.ARROW_FORWARD_IOS, size=14, color=colors["TEXT_MUTED"]),
                    ],
                    alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                    spacing=12,
                ),
                padding=12,
                border_radius=12,
                bgcolor=colors["BG_CARD"],
                border=_border_all(1, colors["BORDER_COLOR"]),
                on_click=on_click,
                ink=True,
            )

        fio_value = " ".join(
            [
                (state.get("last_name") or "").strip(),
                (state.get("first_name") or "").strip(),
                (state.get("middle_name") or "").strip(),
            ]
        ).strip()
        base_info = [
            ft.Text("Профіль користувача", size=22, weight="bold", color=colors["TEXT_COLOR"]),
            ft.Container(height=10),
            ft.Text(f"Поточна роль: {user_role}", size=18, color=colors["ACCENT"]),
            ft.Container(height=6),
            ft.Text(
                f"ПІБ: {fio_value or '-'}",
                size=15,
                color=colors["TEXT_COLOR"],
            ),
            ft.Text(
                f"Email: {(state.get('email') or '').strip() or '-'}",
                size=15,
                color=colors["TEXT_COLOR"],
            ),
            ft.Container(height=20),
        ]

        if user_role == "Адміністратор":
            def open_users(e):
                role_dd = ft.Dropdown(
                    value=state.get("user_role", "Студент"),
                    options=[
                        ft.dropdown.Option("Студент"),
                        ft.dropdown.Option("Викладач"),
                        ft.dropdown.Option("Адміністратор"),
                    ],
                    width=260,
                    label="Встановити роль користувачу",
                )

                def set_role_click(ev):
                    new_role = role_dd.value or "Студент"
                    save_user_role_to_db(new_role)
                    state["user_role"] = new_role
                    telegram_bridge.notify(f"🛡 Роль у додатку змінено на: {new_role}")
                    d.open = False
                    page.update()
                    update_content()
                    show_snack(page, f"Роль змінено на: {new_role}")

                d = ft.AlertDialog(
                    title=ft.Text("Керування користувачами"),
                    content=ft.Column(
                        [
                            ft.Text("Видайте роль користувачу. Тільки адміністратор може змінювати ролі.", size=13, color=colors["TEXT_MUTED"]),
                            ft.Container(height=12),
                            ft.Container(
                                content=ft.Column(
                                    [ft.Text("• Поточний користувач", size=13, color=colors["TEXT_COLOR"])],
                                    spacing=4,
                                ),
                                padding=12,
                                bgcolor=colors["BG_CARD"],
                                border_radius=8,
                            ),
                            ft.Container(height=12),
                            role_dd,
                        ],
                        tight=True,
                        scroll=ft.ScrollMode.AUTO,
                    ),
                    actions=[
                        ft.TextButton("Закрити", on_click=lambda e: (setattr(d, "open", False), page.update())),
                        ft.Button("Встановити роль", on_click=set_role_click, bgcolor=colors["ACCENT"], style=ft.ButtonStyle(color=colors["ACCENT_TEXT"])),
                    ],
                )
                page.overlay.append(d)
                d.open = True
                page.update()

            def open_send_notification_admin(e):
                title_field = ft.TextField(label="Заголовок", width=420)
                msg_field = ft.TextField(label="Текст повідомлення студентам", multiline=True, min_lines=3, max_lines=8, width=420)
                selected_groups = {"items": []}
                groups_preview_ref = ft.Ref[ft.Text]()
                message_dialog_ref = {"dlg": None}

                def refresh_groups_preview():
                    if not groups_preview_ref.current:
                        return
                    if not selected_groups["items"]:
                        groups_preview_ref.current.value = "Групи: усі"
                    else:
                        groups_preview_ref.current.value = "Групи: " + ", ".join(selected_groups["items"][:6]) + (
                            f" (+{len(selected_groups['items']) - 6})" if len(selected_groups["items"]) > 6 else ""
                        )
                    groups_preview_ref.current.update()

                def open_groups_picker(ev):
                    available_groups = get_all_available_groups()
                    checks = []
                    refs: dict[str, ft.Ref[ft.Checkbox]] = {}
                    for g in available_groups:
                        ref = ft.Ref[ft.Checkbox]()
                        refs[g] = ref
                        checks.append(ft.Checkbox(ref=ref, label=g, value=g in selected_groups["items"]))

                    def save_groups_click(_):
                        try:
                            selected_groups["items"] = [g for g in available_groups if refs[g].current and refs[g].current.value]
                            gd.open = False
                            if message_dialog_ref["dlg"] is not None:
                                message_dialog_ref["dlg"].open = True
                            msg_field.focused = True
                            page.update()
                            refresh_groups_preview()
                            show_snack(page, "Групи застосовано. Можна продовжувати писати повідомлення.")
                        except Exception:
                            gd.open = False
                            page.update()
                            show_snack(page, "Не вдалося застосувати групи. Спробуйте ще раз.")

                    gd = ft.AlertDialog(
                        title=ft.Text("Вибір груп"),
                        content=ft.Container(content=ft.Column(checks, spacing=6, scroll=ft.ScrollMode.AUTO), width=420, height=320),
                        actions=[
                            ft.TextButton("Скасувати", on_click=lambda _e: (setattr(gd, "open", False), page.update())),
                            ft.Button("Застосувати", on_click=save_groups_click, bgcolor=colors["ACCENT"], style=ft.ButtonStyle(color=colors["ACCENT_TEXT"])),
                        ],
                    )
                    page.overlay.append(gd)
                    gd.open = True
                    page.update()

                def publish_click(ev):
                    header = (title_field.value or "").strip()
                    text = (msg_field.value or "").strip()
                    target_groups = list(selected_groups["items"])
                    d.open = False
                    page.update()
                    if text or header:
                        author_name = current_user_fio() or "Невідомий користувач"
                        save_announcement(text, title=header, target_groups=target_groups, author_name=author_name, author_role=user_role)
                        telegram_bridge.notify(
                            "\n".join(
                                [
                                    f"📢 {header}" if header else "📢 Оголошення",
                                    text or "",
                                    f"Групи: {', '.join(target_groups)}" if target_groups else "Групи: усі",
                                    f"Автор: {author_name}",
                                ]
                            ).strip()
                        )
                        update_content()
                        show_snack(page, "Опубліковано.")
                    else:
                        show_snack(page, "Введіть заголовок або текст.")

                d = ft.AlertDialog(
                    title=ft.Text("Опублікувати повідомлення"),
                    content=ft.Column(
                        [
                            title_field,
                            msg_field,
                            ft.Row(
                                [
                                    ft.Button("Парсинг груп", icon=ft.icons.SYNC, on_click=lambda _e: trigger_full_groups_parse(show_feedback=True), bgcolor=colors["BORDER_COLOR"]),
                                    ft.Button("Вибрати групу", icon=ft.icons.GROUPS, on_click=open_groups_picker, bgcolor=colors["BORDER_COLOR"]),
                                    ft.Text("Групи: усі", ref=groups_preview_ref, size=12, color=colors["TEXT_MUTED"], expand=True),
                                ],
                                spacing=10,
                            ),
                        ],
                        tight=True,
                        spacing=8,
                    ),
                    actions=[
                        ft.TextButton("Скасувати", on_click=lambda _e: (setattr(d, "open", False), page.update())),
                        ft.Button("Опублікувати", on_click=publish_click, bgcolor=colors["ACCENT"], style=ft.ButtonStyle(color=colors["ACCENT_TEXT"])),
                    ],
                )
                message_dialog_ref["dlg"] = d
                page.overlay.append(d)
                d.open = True
                page.update()
                refresh_groups_preview()

            role_controls = [
                ft.Text("🔧 Панель адміністратора", size=16, weight="bold", color=colors["TEXT_COLOR"]),
                ft.Container(height=8),
                make_action_row(ft.icons.PEOPLE, "Керування користувачами", open_users),
                make_action_row(ft.icons.NOTIFICATIONS_ACTIVE, "Надсилання сповіщень студентам", open_send_notification_admin),
            ]
        elif user_role == "Викладач":
            def open_send_notification(e):
                title_field = ft.TextField(label="Заголовок", width=420)
                msg_field = ft.TextField(label="Текст повідомлення студентам", multiline=True, min_lines=3, max_lines=8, width=420)
                selected_groups = {"items": []}
                groups_preview_ref = ft.Ref[ft.Text]()
                selected_attachment = {"name": "", "path": ""}
                attachment_label_ref = ft.Ref[ft.Text]()
                message_dialog_ref = {"dlg": None}

                def on_file_result(ev: ft.FilePickerResultEvent):
                    if ev.files:
                        selected_attachment["name"] = ev.files[0].name or ""
                        selected_attachment["path"] = ev.files[0].path or ""
                    else:
                        selected_attachment["name"] = ""
                        selected_attachment["path"] = ""
                    if attachment_label_ref.current:
                        attachment_label_ref.current.value = selected_attachment["name"] or "Файл не обрано"
                        attachment_label_ref.current.update()

                file_picker = ft.FilePicker(on_result=on_file_result)
                page.overlay.append(file_picker)

                def pick_file(ev):
                    file_picker.pick_files(allow_multiple=False, dialog_title="Оберіть файл до повідомлення")

                def refresh_groups_preview():
                    if not groups_preview_ref.current:
                        return
                    if not selected_groups["items"]:
                        groups_preview_ref.current.value = "Групи: усі"
                    else:
                        groups_preview_ref.current.value = "Групи: " + ", ".join(selected_groups["items"][:6]) + (
                            f" (+{len(selected_groups['items']) - 6})" if len(selected_groups["items"]) > 6 else ""
                        )
                    groups_preview_ref.current.update()

                def open_groups_picker(ev):
                    available_groups = get_all_available_groups()
                    checks = []
                    refs: dict[str, ft.Ref[ft.Checkbox]] = {}
                    for g in available_groups:
                        ref = ft.Ref[ft.Checkbox]()
                        refs[g] = ref
                        checks.append(
                            ft.Checkbox(
                                ref=ref,
                                label=g,
                                value=g in selected_groups["items"],
                            )
                        )

                    def select_all_click(_):
                        selected_groups["items"] = list(available_groups)
                        for g in available_groups:
                            if refs[g].current:
                                refs[g].current.value = True
                                refs[g].current.update()

                    def clear_all_click(_):
                        selected_groups["items"] = []
                        for g in available_groups:
                            if refs[g].current:
                                refs[g].current.value = False
                                refs[g].current.update()

                    def save_groups_click(_):
                        try:
                            picked = []
                            for g in available_groups:
                                if refs[g].current and refs[g].current.value:
                                    picked.append(g)
                            selected_groups["items"] = picked
                            gd.open = False
                            if message_dialog_ref["dlg"] is not None:
                                message_dialog_ref["dlg"].open = True
                            msg_field.focused = True
                            page.update()
                            refresh_groups_preview()
                            show_snack(page, "Групи застосовано. Можна продовжувати писати повідомлення.")
                        except Exception:
                            gd.open = False
                            page.update()
                            show_snack(page, "Не вдалося застосувати групи. Спробуйте ще раз.")

                    gd = ft.AlertDialog(
                        title=ft.Text("Вибір груп"),
                        content=ft.Container(
                            content=ft.Column(checks, spacing=6, scroll=ft.ScrollMode.AUTO),
                            width=420,
                            height=320,
                        ),
                        actions=[
                            ft.TextButton("Усі", on_click=select_all_click),
                            ft.TextButton("Очистити", on_click=clear_all_click),
                            ft.TextButton("Скасувати", on_click=lambda _e: (setattr(gd, "open", False), page.update())),
                            ft.Button("Застосувати", on_click=save_groups_click, bgcolor=colors["ACCENT"], style=ft.ButtonStyle(color=colors["ACCENT_TEXT"])),
                        ],
                    )
                    page.overlay.append(gd)
                    gd.open = True
                    page.update()

                def publish_click(ev):
                    header = (title_field.value or "").strip()
                    text = (msg_field.value or "").strip()
                    target_groups = list(selected_groups["items"])
                    d.open = False
                    page.update()
                    if text or header:
                        author_name = current_user_fio() or "Невідомий користувач"
                        save_announcement(
                            text,
                            title=header,
                            target_groups=target_groups,
                            author_name=author_name,
                            author_role=user_role,
                            attachment_name=selected_attachment["name"],
                            attachment_path=selected_attachment["path"],
                        )
                        bot_msg = []
                        if header:
                            bot_msg.append(f"📢 {header}")
                        if text:
                            bot_msg.append(text)
                        if target_groups:
                            bot_msg.append(f"Групи: {', '.join(target_groups)}")
                        bot_msg.append(f"Автор: {author_name}")
                        if selected_attachment["name"]:
                            bot_msg.append(f"Файл: {selected_attachment['name']}")
                        telegram_bridge.notify("\n".join(bot_msg))
                        update_content()
                        show_snack(page, "Опубліковано. Студенти побачать повідомлення у сповіщеннях при відкритті додатку.")
                    else:
                        show_snack(page, "Введіть заголовок або текст повідомлення.")

                d = ft.AlertDialog(
                    title=ft.Text("Опублікувати повідомлення"),
                    content=ft.Column(
                        [
                            title_field,
                            msg_field,
                            ft.Row(
                                [
                                    ft.Button(
                                        "Парсинг груп",
                                        icon=ft.icons.SYNC,
                                        on_click=lambda _e: trigger_full_groups_parse(show_feedback=True),
                                        bgcolor=colors["BORDER_COLOR"],
                                    ),
                                    ft.Button(
                                        "Вибрати групу",
                                        icon=ft.icons.GROUPS,
                                        on_click=open_groups_picker,
                                        bgcolor=colors["BORDER_COLOR"],
                                    ),
                                    ft.Text("Групи: усі", ref=groups_preview_ref, size=12, color=colors["TEXT_MUTED"], expand=True),
                                ],
                                spacing=10,
                                alignment=ft.MainAxisAlignment.START,
                            ),
                            ft.Row(
                                [
                                    ft.Button(
                                        "Додати файл",
                                        icon=ft.icons.ATTACH_FILE,
                                        on_click=pick_file,
                                        bgcolor=colors["BORDER_COLOR"],
                                    ),
                                    ft.Text("Файл не обрано", ref=attachment_label_ref, size=12, color=colors["TEXT_MUTED"], expand=True),
                                ],
                                spacing=10,
                                alignment=ft.MainAxisAlignment.START,
                            ),
                            ft.Text(
                                "Після натискання «Опублікувати» повідомлення зʼявиться у сповіщеннях у тих, "
                                "хто відкриє додаток (для обраної групи або для всіх, якщо групу не задано).",
                                size=12,
                                color=colors["TEXT_MUTED"],
                            ),
                        ],
                        tight=True,
                        spacing=8,
                    ),
                    actions=[
                        ft.TextButton("Скасувати", on_click=lambda e: (setattr(d, "open", False), page.update())),
                        ft.Button("Опублікувати", on_click=publish_click, bgcolor=colors["ACCENT"], style=ft.ButtonStyle(color=colors["ACCENT_TEXT"])),
                    ],
                )
                message_dialog_ref["dlg"] = d
                page.overlay.append(d)
                d.open = True
                page.update()
                refresh_groups_preview()

            role_controls = [
                ft.Text("📚 Панель викладача", size=16, weight="bold", color=colors["TEXT_COLOR"]),
                ft.Container(height=8),
                make_action_row(ft.icons.SCHEDULE, "Перегляд та завантаження розкладу", lambda e: go_to_tab(1)),
                make_action_row(ft.icons.NOTIFICATIONS_ACTIVE, "Надсилання сповіщень студентам", open_send_notification),
            ]
        else:
            role_controls = [
                ft.Text("🎓 Панель студента", size=16, weight="bold", color=colors["TEXT_COLOR"]),
                ft.Container(height=8),
                make_action_row(ft.icons.SCHEDULE, "Перегляд розкладу", lambda e: go_to_tab(1)),
                make_action_row(ft.icons.NOTIFICATIONS, "Отримання сповіщень про пари", lambda e: go_to_tab(3)),
            ]

        return ft.Container(
            padding=24,
            content=ft.Column(
                base_info + role_controls,
                spacing=10,
                alignment=ft.MainAxisAlignment.START,
            ),
        )

    def render_auth_gate():
        first_name_ref = ft.Ref[ft.TextField]()
        last_name_ref = ft.Ref[ft.TextField]()
        middle_name_ref = ft.Ref[ft.TextField]()
        email_ref = ft.Ref[ft.TextField]()
        desired_role_ref = ft.Ref[ft.Dropdown]()

        def save_profile_and_continue(e):
            first_name = (first_name_ref.current.value if first_name_ref.current else "").strip()
            last_name = (last_name_ref.current.value if last_name_ref.current else "").strip()
            middle_name = (middle_name_ref.current.value if middle_name_ref.current else "").strip()
            email = (email_ref.current.value if email_ref.current else "").strip().lower()
            if not first_name or not last_name or not middle_name:
                show_snack(page, "Заповніть ПІБ повністю.")
                return
            if not email:
                show_snack(page, "Вкажіть email.")
                return
            if not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
                show_snack(page, "Невірний формат email.")
                return
            state["first_name"] = first_name
            state["last_name"] = last_name
            state["middle_name"] = middle_name
            state["email"] = email
            state["auth_required"] = False
            state["tg_auth_verified"] = False
            requested_role = (desired_role_ref.current.value if desired_role_ref.current else "Викладач").strip() or "Викладач"
            final_role = "Адміністратор" if email == DEFAULT_ADMIN_EMAIL else "Студент"
            state["user_role"] = final_role
            save_user_role_to_db(final_role)
            upsert_user_profile(email, last_name, first_name, middle_name, final_role)
            if email != DEFAULT_ADMIN_EMAIL:
                create_role_change_request(
                    email=email,
                    last_name=last_name,
                    first_name=first_name,
                    middle_name=middle_name,
                    current_role=final_role,
                    requested_role=requested_role,
                )
            save_telegram_auth_profile(
                auth_required=False,
                verified=False,
                first_name=first_name,
                last_name=last_name,
                middle_name=middle_name,
                email=email,
                login_code="",
                login_code_expires_at="",
            )
            page.clean()
            page.add(ft.Row([sidebar, content_area], expand=True))
            state["ui_ready"] = True
            update_content()
            if email == DEFAULT_ADMIN_EMAIL:
                show_snack(page, "Профіль збережено. Вам видано роль Адміністратора для цього пристрою.")
            else:
                show_snack(page, "Профіль збережено. Заявка на зміну ролі відправлена адміну в Telegram.")

        page.add(
            ft.Container(
                expand=True,
                padding=24,
                content=ft.Column(
                    [
                        ft.Text("Перший запуск", size=26, weight="bold", color=colors["TEXT_COLOR"]),
                        ft.Text("Заповніть, будь ласка, прізвище, ім'я та по батькові.", size=14, color=colors["TEXT_MUTED"]),
                        ft.Container(height=12),
                        ft.TextField(ref=last_name_ref, label="Прізвище", value=state.get("last_name", ""), width=420),
                        ft.TextField(ref=first_name_ref, label="Ім'я", value=state.get("first_name", ""), width=420),
                        ft.TextField(ref=middle_name_ref, label="По батькові", value=state.get("middle_name", ""), width=420),
                        ft.TextField(ref=email_ref, label="Email", value=state.get("email", ""), width=420),
                        ft.Dropdown(
                            ref=desired_role_ref,
                            label="Бажана роль (заявка)",
                            value="Викладач",
                            width=420,
                            options=[
                                ft.dropdown.Option("Викладач"),
                                ft.dropdown.Option("Адміністратор"),
                            ],
                        ),
                        ft.Button("Продовжити", on_click=save_profile_and_continue, bgcolor=colors["ACCENT"], style=ft.ButtonStyle(color=colors["ACCENT_TEXT"])),
                    ],
                    spacing=10,
                    horizontal_alignment=ft.CrossAxisAlignment.START,
                ),
            )
        )

    if state.get("auth_required", True):
        state["ui_ready"] = False
        render_auth_gate()
    else:
        page.add(
            ft.Row(
                [
                    sidebar,
                    content_area,
                ],
                expand=True,
            )
        )
        state["ui_ready"] = True
        update_content()
        unread_for_bot = get_unread_announcements(include_meta=True)
        if unread_for_bot and state.get("telegram_enabled"):
            for _, msg, _, _ in unread_for_bot[:5]:
                if (msg or "").strip():
                    telegram_bridge.notify(f"📢 Нове повідомлення:\n{msg.strip()}")
            if len(unread_for_bot) > 5:
                telegram_bridge.notify(f"Ще {len(unread_for_bot) - 5} повідомлень доступно в додатку.")
        show_pending_announcements(page, state["notifications_enabled"], target_group=current_target_group())

    async def preload_schedule():
        state["schedule_error"] = ""
        try:
            await asyncio.sleep(1.0)
            p = await asyncio.to_thread(get_parser)
            if p is None:
                state["schedule_error"] = "Не вдалося підключитися до lib.istu.edu.ua. Перевірте інтернет."
                state["schedule_files"] = []
            else:
                state["schedule_url"] = p.base_url + DEFAULT_SCHEDULE_PATH
                state["schedule_root_url"] = p.base_url + DEFAULT_SCHEDULE_PATH
                try:
                    files = await asyncio.to_thread(p.parse_files, state["schedule_url"])
                    if not files:
                        await asyncio.sleep(0.25)
                        files = await asyncio.to_thread(p.parse_files, state["schedule_url"])
                    if not files:
                        reset_parser()
                        p2 = await asyncio.to_thread(get_parser)
                        if p2:
                            state["schedule_url"] = p2.base_url + DEFAULT_SCHEDULE_PATH
                            state["schedule_root_url"] = p2.base_url + DEFAULT_SCHEDULE_PATH
                            files = await asyncio.to_thread(p2.parse_files, state["schedule_url"])
                    if files:
                        state["schedule_files"] = files
                except (requests.exceptions.ConnectionError, OSError):
                    state["schedule_error"] = "Помилка з'єднання. Перевірте інтернет або VPN."
                    state["schedule_files"] = []
                except requests.exceptions.Timeout:
                    state["schedule_error"] = "Час очікування вийшов. Сайт не відповідає."
                    state["schedule_files"] = []
                except Exception as e:
                    state["schedule_error"] = f"Помилка: {type(e).__name__}"
                    state["schedule_files"] = []
            update_content()
        except Exception:
            state["schedule_error"] = "Не вдалося завантажити розклад."
            state["schedule_files"] = []
            update_content()

    page.run_task(preload_schedule)

    async def preload_all_groups():
        await asyncio.sleep(1.2)
        try:
            groups = await asyncio.to_thread(collect_all_groups_from_schedule_tree, state.get("schedule_root_url", ""))
            if groups:
                merged = sorted(set(normalize_group_names((state.get("all_schedule_groups") or []) + groups)))
                state["all_schedule_groups"] = merged
        except Exception:
            pass

    page.run_task(preload_all_groups)
    telegram_bridge.start()

    async def check_new_announcements_on_start():
        await asyncio.sleep(1.0)
        unread = get_unread_announcements(target_group=current_target_group())
        if unread:
            count = len(unread)
            message = f"Є нові повідомлення ({count}). Щоб подивитись, перейдіть в налаштування."
            show_snack(page, message)
            show_system_notification(
                "МНТУ Помічник — нові повідомлення",
                message,
                timeout=10
            )

    page.run_task(check_new_announcements_on_start)

    async def periodic_check():
        print("[NOTIFICATION] Запущено фоновий перевірку уведомлень")
        try:

            await asyncio.sleep(2)
            print("[NOTIFICATION] Перша перевірка уведомлень...")
            try:
                notifications = find_upcoming_entries()
                print(f"[NOTIFICATION] Перша перевірка: знайдено {len(notifications)} уведомлень")
                for notif_data in notifications:
                    show_class_notification(page, notif_data, colors)
            except Exception as e:
                print(f"[ERROR] Помилка при першій перевірці: {e}")
                import traceback
                traceback.print_exc()

            iteration = 0
            while True:
                await asyncio.sleep(10)
                iteration += 1
                try:
                    if iteration % 6 == 0:
                        print(f"[NOTIFICATION] Перевірка {iteration}...")
                    notifications = find_upcoming_entries()
                    if notifications:
                        print(f"[NOTIFICATION] Знайдено {len(notifications)} нових уведомлень!")
                        try:
                            update_content()
                        except Exception as e:
                            print(f"[ERROR] Помилка при оновленні контенту: {e}")
                    for notif_data in notifications:
                        try:
                            show_class_notification(page, notif_data, colors)
                        except Exception as e:
                            print(f"[ERROR] Помилка при показі сповіщення: {e}")
                except Exception as e:
                    print(f"[ERROR] Помилка в periodic_check (ітерація {iteration}): {e}")
                    import traceback
                    traceback.print_exc()

                    continue
        except asyncio.CancelledError:
            print("[INFO]: Розклад закритий")
        except Exception as ex:
            print(f"[ERROR] Критична помилка в periodic_check: {ex}")
            import traceback
            traceback.print_exc()

    page.run_task(periodic_check)
